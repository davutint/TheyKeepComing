from __future__ import annotations

import math
import os
import random
import textwrap
from pathlib import Path

from PIL import Image, ImageDraw, ImageFont

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(r"C:\GithubProjeler\TheyKeepComing\IncremantalDots")
ASSET_DIR = ROOT / "Assets" / "Docs" / "References" / "GameDesignBlueprint"
TMP_DIR = ROOT / "tmp" / "dead_walls_blueprint"
OUT_DOCX = ROOT / "Assets" / "Docs" / "DEAD_WALLS_GAME_DESIGN_BLUEPRINT_v1.0.docx"
COVER = ASSET_DIR / "dead_walls_blueprint_cover.png"

ASSET_DIR.mkdir(parents=True, exist_ok=True)
TMP_DIR.mkdir(parents=True, exist_ok=True)


# Dead Walls Nightfall named override on compact_reference_guide.
COLORS = {
    "night": "0B1726",
    "indigo": "18304E",
    "heart": "B63A48",
    "fire": "E99A32",
    "frost": "4FA9C7",
    "ink": "1B2735",
    "body": "2F3B49",
    "muted": "687789",
    "stone": "D9E0E8",
    "paper": "F6F8FB",
    "pale_blue": "EAF2F8",
    "pale_red": "F9EBED",
    "pale_gold": "FFF4DF",
    "pale_green": "EAF5EF",
    "green": "2F7D5A",
    "white": "FFFFFF",
    "black": "000000",
}


def rgb(hex_value: str) -> RGBColor:
    return RGBColor.from_string(hex_value.replace("#", ""))


ARIAL = r"C:\Windows\Fonts\arial.ttf"
ARIAL_BOLD = r"C:\Windows\Fonts\arialbd.ttf"


def pil_font(size: int, bold: bool = False):
    path = ARIAL_BOLD if bold else ARIAL
    return ImageFont.truetype(path, size=size)


def wrap(draw: ImageDraw.ImageDraw, text: str, font, width: int) -> list[str]:
    words = text.split()
    lines: list[str] = []
    current = ""
    for word in words:
        test = word if not current else current + " " + word
        if draw.textbbox((0, 0), test, font=font)[2] <= width:
            current = test
        else:
            if current:
                lines.append(current)
            current = word
    if current:
        lines.append(current)
    return lines


def draw_centered(draw, box, text, font, fill, max_width=None, line_gap=8):
    x0, y0, x1, y1 = box
    lines = wrap(draw, text, font, max_width or int(x1 - x0 - 24))
    heights = [draw.textbbox((0, 0), line, font=font)[3] for line in lines]
    total = sum(heights) + line_gap * max(0, len(lines) - 1)
    y = y0 + (y1 - y0 - total) / 2
    for line, h in zip(lines, heights):
        bb = draw.textbbox((0, 0), line, font=font)
        x = x0 + (x1 - x0 - (bb[2] - bb[0])) / 2
        draw.text((x, y), line, font=font, fill=fill)
        y += h + line_gap


def rounded_box(draw, box, fill, outline=None, radius=26, width=3):
    draw.rounded_rectangle(box, radius=radius, fill=fill, outline=outline, width=width)


def arrow(draw, start, end, color, width=6):
    draw.line([start, end], fill=color, width=width)
    angle = math.atan2(end[1] - start[1], end[0] - start[0])
    size = 18
    p1 = (end[0] - size * math.cos(angle - 0.5), end[1] - size * math.sin(angle - 0.5))
    p2 = (end[0] - size * math.cos(angle + 0.5), end[1] - size * math.sin(angle + 0.5))
    draw.polygon([end, p1, p2], fill=color)


def save_canvas(img: Image.Image, name: str) -> Path:
    path = ASSET_DIR / name
    img.save(path, dpi=(180, 180))
    return path


def make_core_loop() -> Path:
    img = Image.new("RGB", (2160, 1098), "#F6F8FB")
    d = ImageDraw.Draw(img)
    center = (610, 555)
    outer = (210, 155, 1010, 955)
    inner = (365, 310, 855, 800)
    phases = [
        (30, "GÜNDÜZ", "30 sn", "#E9B55A", "#1B2735"),
        (5, "AKŞAM", "5 sn", "#B96B52", "#FFFFFF"),
        (20, "GECE", "20 sn", "#17304E", "#FFFFFF"),
        (5, "ŞAFAK", "5 sn", "#4FA9C7", "#0B1726"),
    ]
    start = -90.0
    for duration, label, seconds, color, text_color in phases:
        end = start + duration / 60.0 * 360.0
        d.pieslice(outer, start=start, end=end, fill=color, outline="#F6F8FB", width=8)
        mid = math.radians((start + end) / 2)
        x = center[0] + 315 * math.cos(mid)
        y = center[1] + 315 * math.sin(mid)
        draw_centered(d, (x - 125, y - 55, x + 125, y + 55), f"{label}\n{seconds}",
                      pil_font(26, True), text_color, max_width=235, line_gap=4)
        start = end
    d.ellipse(inner, fill="#F6F8FB")
    draw_centered(d, (390, 390, 830, 625), "60 SANİYE\nKESİNTİSİZ KUŞATMA",
                  pil_font(44, True), "#0B1726", max_width=420, line_gap=12)
    cards = [
        (1190, 205, "GÜNDÜZ", "İşçi dağıt • satın al • onar", "#FFF4DF", "#9A681E"),
        (1190, 445, "GECE", "Sürü yoğunluğu zirve yapar", "#EAF2F8", "#17304E"),
        (1190, 685, "ŞAFAK", "Yeni nüfus • daha yüksek taban", "#E7F4F7", "#287C96"),
    ]
    for x, y, head, body, fill, accent in cards:
        rounded_box(d, (x, y, 2070, y + 175), fill, outline=accent, radius=28, width=4)
        d.text((x + 38, y + 30), head, font=pil_font(32, True), fill=accent)
        d.text((x + 38, y + 95), body, font=pil_font(24), fill="#2F3B49")
    return save_canvas(img, "fig_core_loop.png")


def make_spawn_curve() -> Path:
    img = Image.new("RGB", (2160, 1008), "#F6F8FB")
    d = ImageDraw.Draw(img)
    d.text((95, 55), "AYNI DÜŞMAN, SABİT İSTATİSTİKLER, KATLANAN ADET",
           font=pil_font(43, True), fill="#0B1726")
    plot = (210, 195, 2050, 820)
    x0, y0, x1, y1 = plot
    for i in range(6):
        y = y0 + i * (y1 - y0) / 5
        d.line((x0, y, x1, y), fill="#D9E0E8", width=2)
    for i in range(20):
        x = x0 + i * (x1 - x0) / 19
        if i in (0, 4, 9, 14, 19):
            d.text((x - 12, y1 + 24), str(i + 1), font=pil_font(20), fill="#687789")
    d.text((965, 900), "Koşu günü", font=pil_font(24, True), fill="#2F3B49")
    d.text((30, 465), "Yoğunluk\nendeksi", font=pil_font(22, True), fill="#2F3B49")

    day_values = [1.18 ** i for i in range(20)]
    night_values = [v * 5.5 for v in day_values]
    max_log = math.log10(max(night_values))
    def point(index, value):
        x = x0 + index * (x1 - x0) / 19
        y = y1 - (math.log10(value) / max_log) * (y1 - y0)
        return (x, y)
    day_points = [point(i, v) for i, v in enumerate(day_values)]
    night_points = [point(i, v) for i, v in enumerate(night_values)]
    d.polygon(night_points + list(reversed(day_points)), fill="#DCECF2")
    d.line(day_points, fill="#E9A841", width=8, joint="curve")
    d.line(night_points, fill="#18304E", width=10, joint="curve")
    d.line((1440, 130, 1530, 130), fill="#18304E", width=10)
    d.text((1550, 112), "Gece yoğunluğu", font=pil_font(23, True), fill="#18304E")
    d.line((1440, 170, 1530, 170), fill="#E9A841", width=8)
    d.text((1550, 152), "Gündüz yoğunluğu", font=pil_font(23, True), fill="#9A681E")
    rounded_box(d, (245, 245, 675, 355), "#F9EBED", outline="#B63A48", radius=20, width=3)
    d.text((275, 280), "HP / HASAR / HIZ SABİT", font=pil_font(24, True), fill="#B63A48")
    rounded_box(d, (1425, 385, 1950, 500), "#EAF2F8", outline="#18304E", radius=20, width=3)
    d.text((1460, 420), "ZORLUK = SAYI + AKIŞ", font=pil_font(27, True), fill="#18304E")
    return save_canvas(img, "fig_spawn_curve.png")


def make_battlefield() -> Path:
    img = Image.new("RGB", (1800, 860), "#F6F8FB")
    d = ImageDraw.Draw(img)
    title = pil_font(48, True)
    label = pil_font(31, True)
    body = pil_font(24)
    d.text((65, 42), "TEK CEPHELİ SABİT KALE KOMPOZİSYONU", font=title, fill="#0B1726")
    zones = [
        (65, 155, 475, 760, "KALE İÇİ", "Hazır binalar\nTemsilî işçiler\nSayısal nüfus", "#E7C487"),
        (500, 155, 805, 760, "DUVAR", "Tek ortak HP\n40 okçu tile'ı\nOnarım gündüz", "#B7C2CE"),
        (830, 155, 1735, 760, "KUŞATMA ALANI", "Sağdan gelen tek prefab\nÇoklu spawn noktası\n10.000 aktif düşman hedefi", "#B8CDDC"),
    ]
    for x0, y0, x1, y1, head, txt, fill in zones:
        rounded_box(d, (x0, y0, x1, y1), fill, outline="#7A8796", radius=30, width=4)
        d.text((x0 + 28, y0 + 28), head, font=label, fill="#0B1726")
        for i, line in enumerate(txt.split("\n")):
            d.text((x0 + 28, y0 + 95 + i * 45), line, font=body, fill="#2F3B49")
    # wall and archers
    d.rectangle((635, 345, 675, 670), fill="#556270")
    for i in range(10):
        x = 560 + (i % 3) * 34
        y = 360 + (i // 3) * 65
        d.ellipse((x, y, x + 22, y + 22), fill="#18304E")
    # enemy stream
    rng = random.Random(44)
    for _ in range(280):
        x = rng.randint(950, 1685)
        y = rng.randint(330, 705)
        r = rng.randint(3, 7)
        d.ellipse((x-r, y-r, x+r, y+r), fill="#263D55")
    arrow(d, (1685, 270), (765, 270), "#B63A48", 9)
    d.text((1090, 215), "KESİNTİSİZ SAĞDAN SOLA BASKI", font=body, fill="#B63A48")
    return save_canvas(img, "fig_battlefield.png")


def make_economy_flow() -> Path:
    img = Image.new("RGB", (1800, 940), "#F6F8FB")
    d = ImageDraw.Draw(img)
    title = pil_font(48, True)
    head = pil_font(28, True)
    body = pil_font(22)
    d.text((65, 42), "KAYNAK ÜRETİMİ VE TEK SEFERLİK HARCAMA", font=title, fill="#0B1726")
    sources = [
        ("FARM", "FOOD", "#E5B65B"), ("LUMBERYARD", "WOOD", "#A8794E"),
        ("QUARRY", "STONE", "#8796A6"), ("MINE", "IRON", "#50677E")
    ]
    targets = [
        ("Yeni nüfus", "Food bir kez"), ("Yatak", "artan maliyet"),
        ("Okçu", "anında satın alma"), ("Ok", "Wood ile anında"),
        ("Onarım", "gündüz Stone"), ("Kapasite", "tekrar yükselt")
    ]
    for i, (building, resource, color) in enumerate(sources):
        y = 175 + i * 170
        rounded_box(d, (70, y, 485, y + 112), color, outline="#6D7782", radius=22, width=3)
        d.text((95, y + 18), building, font=head, fill="#FFFFFF")
        d.text((95, y + 62), resource, font=body, fill="#FFFFFF")
        arrow(d, (500, y + 56), (760, y + 56), "#687789", 6)
    rounded_box(d, (770, 190, 1080, 790), "#EAF2F8", outline="#4FA9C7", radius=28, width=4)
    draw_centered(d, (790, 210, 1060, 770), "İŞÇİ SAYISI\n×\nİŞÇİ VERİMLİLİĞİ\n=\nÜRETİM", pil_font(31, True), "#18304E", 250, 14)
    for i, (name, desc) in enumerate(targets):
        col, row = i % 2, i // 2
        x = 1180 + col * 285
        y = 175 + row * 225
        rounded_box(d, (x, y, x + 245, y + 155), "#FFFFFF", outline="#B9C5D0", radius=20, width=3)
        draw_centered(d, (x + 10, y + 15, x + 235, y + 85), name, head, "#0B1726", 215)
        draw_centered(d, (x + 14, y + 86, x + 231, y + 145), desc, body, "#687789", 210)
        arrow(d, (1088, 490), (x - 14, y + 75), "#B63A48", 5)
    d.text((725, 842), "PASİF GİDER YOK  |  KAYNAK SADECE OYUNCU SATIN ALDIĞINDA AZALIR",
           font=pil_font(25, True), fill="#2F7D5A")
    return save_canvas(img, "fig_economy_flow.png")


def random_points_in_diamond(seed: int, count: int, cx: int, cy: int, rx: int, ry: int, min_dist=15):
    rng = random.Random(seed)
    pts = []
    attempts = 0
    while len(pts) < count and attempts < 10000:
        attempts += 1
        x = rng.uniform(cx-rx, cx+rx)
        y = rng.uniform(cy-ry, cy+ry)
        if abs(x-cx)/rx + abs(y-cy)/ry > 0.82:
            continue
        if any((x-px)**2 + (y-py)**2 < min_dist**2 for px, py in pts):
            continue
        pts.append((x, y))
    return pts


def make_archer_placement() -> Path:
    img = Image.new("RGB", (1800, 930), "#F6F8FB")
    d = ImageDraw.Draw(img)
    title = pil_font(48, True)
    head = pil_font(28, True)
    body = pil_font(22)
    d.text((65, 42), "40 TILE × 25 NOKTA = 1.000 GÖRÜNÜR OKÇU", font=title, fill="#0B1726")
    d.text((75, 118), "Mevcut outside tilemap sırası korunur; nokta deseni her tile için sabit seed ile farklıdır.", font=body, fill="#687789")
    # overview line of 40 diamonds
    for i in range(40):
        x = 85 + i * 40
        cy = 270 + int(55 * math.sin(i / 39 * math.pi))
        diamond = [(x, cy), (x+22, cy-12), (x+44, cy), (x+22, cy+12)]
        d.polygon(diamond, fill="#DCE5ED", outline="#7D8B99")
        if i < 8:
            d.ellipse((x+19, cy-3, x+25, cy+3), fill="#B63A48")
    d.text((75, 345), "DOLDURMA: önce 40 tile'a birer okçu, sonra her tile'ın 2. noktası...", font=head, fill="#18304E")
    # zoom tile
    cx, cy, rx, ry = 480, 640, 330, 190
    d.polygon([(cx,cy-ry),(cx+rx,cy),(cx,cy+ry),(cx-rx,cy)], fill="#EAF2F8", outline="#4FA9C7")
    pts = random_points_in_diamond(101, 25, cx, cy, rx, ry, 45)
    for idx, (x,y) in enumerate(pts, 1):
        d.ellipse((x-11,y-11,x+11,y+11), fill="#18304E", outline="#FFFFFF", width=2)
        d.text((x+13,y-12), str(idx), font=pil_font(14, True), fill="#687789")
    rounded_box(d, (930, 450, 1710, 825), "#FFFFFF", outline="#B9C5D0", radius=24, width=3)
    notes = [
        "Noktalar elmas sınırın içinde kalır.",
        "Minimum mesafe üst üste binmeyi azaltır.",
        "Her tile farklı ama tekrar üretilebilir desene sahiptir.",
        "Kayıt yüklenince formasyon değişmez.",
        "Okçu türü konumu belirlemez; satın alma sırası belirler.",
    ]
    d.text((970, 485), "YERLEŞİM SÖZLEŞMESİ", font=head, fill="#B63A48")
    for i, note in enumerate(notes):
        y = 555 + i*52
        d.ellipse((972,y+7,985,y+20), fill="#E99A32")
        d.text((1002,y), note, font=body, fill="#2F3B49")
    return save_canvas(img, "fig_archer_placement.png")


def make_run_flow() -> Path:
    img = Image.new("RGB", (1800, 880), "#F6F8FB")
    d = ImageDraw.Draw(img)
    title = pil_font(48, True)
    head = pil_font(28, True)
    body = pil_font(21)
    d.text((65, 42), "KOŞU, MENÜ VE ÖLÜM SÖZLEŞMESİ", font=title, fill="#0B1726")
    boxes = {
        "run": (70, 260, 430, 520, "AKTİF KOŞU", "60 sn döngü\nKaynaklar\nCastle Heart graph"),
        "menu": (610, 115, 1010, 350, "ANA MENÜ", "Koşu donar\nAuto-save\nDEVAM ET"),
        "death": (610, 500, 1010, 735, "WALL HP = 0", "Anında Game Over\nÇöküş aşaması yok\nKoşu kapanır"),
        "meta": (1190, 500, 1690, 735, "META EKRANI", "Kalıcı para\nSade upgrade listesi\nYeni koşu"),
    }
    fills = {"run":"#EAF2F8","menu":"#FFF4DF","death":"#F9EBED","meta":"#EAF5EF"}
    outlines = {"run":"#4FA9C7","menu":"#E99A32","death":"#B63A48","meta":"#2F7D5A"}
    for key,(x0,y0,x1,y1,h,t) in boxes.items():
        rounded_box(d,(x0,y0,x1,y1),fills[key],outline=outlines[key],radius=28,width=4)
        draw_centered(d,(x0+20,y0+20,x1-20,y0+90),h,head,"#0B1726")
        draw_centered(d,(x0+25,y0+95,x1-25,y1-25),t,body,"#2F3B49")
    arrow(d,(430,330),(610,235),"#687789",7)
    d.text((455,245),"Çık",font=body,fill="#687789")
    arrow(d,(610,300),(430,420),"#2F7D5A",7)
    d.text((465,380),"Devam",font=body,fill="#2F7D5A")
    arrow(d,(430,475),(610,610),"#B63A48",7)
    d.text((455,535),"Duvar biter",font=body,fill="#B63A48")
    arrow(d,(1010,620),(1190,620),"#2F7D5A",7)
    arrow(d,(1440,500),(430,465),"#2F7D5A",7)
    d.text((1050,420),"Yeni koşu: run state sıfır, meta kalır",font=head,fill="#2F7D5A")
    d.text((70,790),"Gönüllü reset yok  |  Offline ilerleme yok  |  Tek aktif koşu  |  Eski save noktasına dönüş yok",font=pil_font(25,True),fill="#18304E")
    return save_canvas(img, "fig_run_flow.png")


def make_tech_graph() -> Path:
    img = Image.new("RGB", (1800, 1000), "#F6F8FB")
    d = ImageDraw.Draw(img)
    title = pil_font(44, True)
    node_font = pil_font(19, True)
    small = pil_font(16)
    d.text((65, 38), "PROCEDURAL CASTLE HEART - ÖRNEK KOŞU, SABİT AĞAÇ DEĞİL", font=title, fill="#0B1726")
    center = (900, 505)
    dirs = {
        "ORDU": ((1450, 505), "#B63A48"),
        "SAVUNMA": ((900, 145), "#6B7B8D"),
        "ÜRETİM": ((350, 505), "#E99A32"),
        "HEART / BÜYÜ": ((900, 865), "#4FA9C7"),
    }
    # Branch veins
    for label,(pos,color) in dirs.items():
        arrow(d, center, pos, color, 9)
        d.text((pos[0]-80,pos[1]-64),label,font=small,fill=color)
    rounded_box(d,(790,425,1010,585),"#F9EBED",outline="#B63A48",radius=55,width=6)
    draw_centered(d,(800,438,1000,570),"CASTLE\nHEART",pil_font(27,True),"#B63A48")
    nodes = [
        (1250,430,"Rapid kilidi","#F9EBED"),(1510,365,"Atış hızı ∞","#F9EBED"),(1510,545,"Nadir evolution","#FFF4DF"),
        (790,210,"Wall HP ∞","#EEF1F4"),(1120,110,"Onarım verimi","#EEF1F4"),(665,95,"Keystone","#FFF4DF"),
        (510,420,"İşçi verimi ∞","#FFF4DF"),(210,330,"Kapasite","#FFF4DF"),(210,555,"Gizli node","#E9EDF1"),
        (790,730,"Fireball","#EAF2F8"),(1140,830,"Patlama alanı ∞","#EAF2F8"),(635,875,"Rally","#EAF2F8")
    ]
    for x,y,text,fill in nodes:
        outline = "#B9C5D0" if text != "Gizli node" else "#9FAAB6"
        rounded_box(d,(x-110,y-45,x+110,y+45),fill,outline=outline,radius=24,width=3)
        draw_centered(d,(x-102,y-38,x+102,y+38),text,node_font,"#1B2735",195)
    # local connections
    for a,b,color in [((1010,485),(1140,450),"#B63A48"),((1360,430),(1400,385),"#B63A48"),((1360,460),(1400,545),"#B63A48"),
                      ((900,425),(900,255),"#6B7B8D"),((1000,190),(1010,135),"#6B7B8D"),((805,185),(735,130),"#6B7B8D"),
                      ((790,505),(620,465),"#E99A32"),((400,425),(315,365),"#E99A32"),((400,500),(315,555),"#E99A32"),
                      ((900,585),(900,685),"#4FA9C7"),((1000,750),(1040,805),"#4FA9C7"),((800,760),(725,840),"#4FA9C7")]:
        arrow(d,a,b,color,5)
    d.text((65, 940), "Sabit: yönlerin kimliği ve temel sistemlerin bulunması  |  Değişken: node, bağlantı, sıra, nadirlik, Keystone", font=pil_font(21,True), fill="#2F3B49")
    return save_canvas(img, "fig_tech_graph.png")


def make_hud_wireframe() -> Path:
    img = Image.new("RGB", (1800, 1010), "#0B1726")
    d = ImageDraw.Draw(img)
    title = pil_font(42, True)
    label = pil_font(22, True)
    small = pil_font(17)
    d.text((55, 32), "SABİT KAMERA + MİNİMAL HUD YERLEŞİMİ", font=title, fill="#FFFFFF")
    # world frame
    rounded_box(d,(55,105,1745,945),"#14263A",outline="#38506A",radius=24,width=4)
    # castle and horde abstract
    d.rectangle((150,380,510,820),fill="#31465A")
    d.rectangle((520,310,565,830),fill="#738191")
    rng=random.Random(9)
    for _ in range(240):
        x=rng.randint(680,1660); y=rng.randint(300,850); r=rng.randint(2,5)
        d.ellipse((x-r,y-r,x+r,y+r),fill="#55708A")
    # top resources
    rounded_box(d,(95,135,1010,195),"#0F1E30",outline="#4D637A",radius=16,width=2)
    draw_centered(d,(110,140,995,190),"WOOD  |  STONE  |  IRON  |  FOOD  |  ARROWS  |  ESSENCE  |  POP",small,"#FFFFFF")
    rounded_box(d,(750,215,1050,260),"#0F1E30",outline="#E99A32",radius=15,width=2)
    draw_centered(d,(760,220,1040,255),"FAZ GÖSTERİMİ - POLISH OPEN",small,"#E9B55A")
    # wall bar
    rounded_box(d,(150,835,600,885),"#0F1E30",outline="#B63A48",radius=14,width=2)
    d.rectangle((170,852,465,870),fill="#B63A48")
    d.text((485,842),"WALL 73%",font=small,fill="#FFFFFF")
    # bottom buttons
    rounded_box(d,(95,870,455,925),"#0F1E30",outline="#4FA9C7",radius=15,width=2)
    draw_centered(d,(105,875,445,920),"WORKERS  |  HOUSING",small,"#FFFFFF")
    rounded_box(d,(650,830,1150,925),"#0F1E30",outline="#E99A32",radius=18,width=3)
    draw_centered(d,(670,840,1130,915),"[1] FIREBALL   [2] RALLY   [3] REPAIR",label,"#FFFFFF")
    rounded_box(d,(1345,870,1705,925),"#0F1E30",outline="#B63A48",radius=15,width=2)
    draw_centered(d,(1355,875,1695,920),"ARCHERS  |  CASTLE HEART",small,"#FFFFFF")
    d.text((92,960),"Pan yok  |  Zoom yok  |  Gate/Core barları yok  |  Düşman tahmini yok  |  Heart açılınca oyun durur",font=pil_font(20,True),fill="#B8CDDC")
    return save_canvas(img, "fig_hud_wireframe.png")


def make_architecture() -> Path:
    img = Image.new("RGB", (1800, 980), "#F6F8FB")
    d=ImageDraw.Draw(img)
    title=pil_font(44,True); head=pil_font(24,True); body=pil_font(18)
    d.text((65,40),"VERİ ODAKLI UYGULAMA HARİTASI",font=title,fill="#0B1726")
    cols=[(65,250,450,"TANIM ASSETLERİ","Enemy Definition\nArcher Definition\nTech Node Pool\nMeta Upgrade\nDifficulty Profile","#FFF4DF","#E99A32"),
          (535,170,975,"RUN RUNTIME","Cycle + Spawn Budget\nResources + Population\nArcher Counts + Ammo\nCouncil Schedule + Effects\nGenerated Heart Graph\nWall HP + Cooldowns","#EAF2F8","#4FA9C7"),
          (1060,250,1435,"SUNUM","World Workers\n10k Horde Render\nArcher Formation\nHUD Drawers\nFireball Targeting","#F9EBED","#B63A48"),
          (1060,610,1435,"KALICILIK","Single Run Save\nGenerated Graph\nCouncil Memory\nMeta State\nTutorial Flags","#EAF5EF","#2F7D5A")]
    for x0,y0,x1,h,txt,fill,out in cols:
        y1=y0+280
        rounded_box(d,(x0,y0,x1,y1),fill,outline=out,radius=25,width=4)
        draw_centered(d,(x0+15,y0+15,x1-15,y0+70),h,head,"#0B1726")
        draw_centered(d,(x0+25,y0+80,x1-25,y1-20),txt,body,"#2F3B49")
    arrow(d,(450,390),(535,390),"#687789",7)
    arrow(d,(975,330),(1060,330),"#687789",7)
    arrow(d,(975,500),(1060,690),"#2F7D5A",7)
    rounded_box(d,(1515,250,1730,890),"#0B1726",outline="#0B1726",radius=24,width=3)
    draw_centered(d,(1530,270,1715,870),"KURAL\n\nYeni düşman prefabı veya yeni node eklemek ana UI / spawn kodunu değiştirmemeli.\n\nLegacy Mobile* isimleri sırf ürün yönü değişti diye yeniden adlandırılmaz.",body,"#FFFFFF",165,10)
    d.text((65,890),"MonoBehaviour: UI, input, visual bridge  |  ECS: sürü, okçu, mermi, durum ve performans kritik simülasyon",font=pil_font(21,True),fill="#18304E")
    return save_canvas(img,"fig_architecture.png")


def make_delivery_sequence() -> Path:
    img=Image.new("RGB",(1800,920),"#F6F8FB"); d=ImageDraw.Draw(img)
    title=pil_font(44,True); head=pil_font(22,True); small=pil_font(17)
    d.text((65,38),"UYGULAMA SIRASI - TAKVİM DEĞİL, BAĞIMLILIK AKIŞI",font=title,fill="#0B1726")
    stages=[
        ("A","Sözleşmeler","Run state\nTek wall\nKaynak kuralları"),
        ("B","Kuşatma","60 sn cycle\nSpawn budget\nTek prefab"),
        ("C","Ekonomi","4 bina\nİşçi oranı\nYatak + nüfus"),
        ("D","Okçular","1.000 cap\n40×25 yerleşim\nAmmo"),
        ("E","Castle Heart","Run graph\nEssence\nUpgrade etkileri"),
        ("F","Council","3/6/9... takvim\nBağlam duyarlı\nCurated içerik"),
        ("G","Yetenekler","Fireball\nRally\nEmergency Repair"),
        ("H","Meta + Save","Game Over\nMeta listesi\nContinue"),
        ("I","Ürün Kapısı","HUD polish\n10k stress\nOnboarding + QA"),
    ]
    for i,(code,name,txt) in enumerate(stages):
        row=0 if i<5 else 1
        col=i if row==0 else 4-(i-5)
        x=65+col*340; y=170+row*315
        fill="#EAF2F8" if row==0 else "#F9EBED"
        out="#4FA9C7" if row==0 else "#B63A48"
        rounded_box(d,(x,y,x+300,y+215),fill,outline=out,radius=24,width=4)
        d.ellipse((x+20,y+20,x+78,y+78),fill=out)
        draw_centered(d,(x+20,y+20,x+78,y+78),code,head,"#FFFFFF")
        d.text((x+95,y+25),name,font=head,fill="#0B1726")
        draw_centered(d,(x+24,y+90,x+276,y+195),txt,small,"#2F3B49")
        if row==0 and i<4:
            arrow(d,(x+300,y+108),(x+335,y+108),"#687789",5)
        elif i==4:
            arrow(d,(x+150,y+215),(x+150,y+305),"#687789",5)
        elif row==1 and i<8:
            arrow(d,(x,y+108),(x-35,y+108),"#687789",5)
    d.text((65,840),"Her paket oynanabilir bir sonuç ve ölçülebilir kabul kapısı üretir. Süre tahmini bu belgeye bilinçli olarak yazılmamıştır.",font=pil_font(21,True),fill="#2F7D5A")
    return save_canvas(img,"fig_delivery_sequence.png")


def create_figures():
    return {
        "core_loop": make_core_loop(),
        "spawn_curve": make_spawn_curve(),
        "battlefield": make_battlefield(),
        "economy": make_economy_flow(),
        "placement": make_archer_placement(),
        "run_flow": make_run_flow(),
        "tech": make_tech_graph(),
        "hud": make_hud_wireframe(),
        "architecture": make_architecture(),
        "delivery": make_delivery_sequence(),
    }


# ----------------------------- DOCX HELPERS -----------------------------

def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=100, start=140, bottom=100, end=140):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_border(cell, **edges):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge_name, cfg in edges.items():
        tag = "start" if edge_name == "left" else "end" if edge_name == "right" else edge_name
        edge = borders.find(qn(f"w:{tag}"))
        if edge is None:
            edge = OxmlElement(f"w:{tag}")
            borders.append(edge)
        edge.set(qn("w:val"), cfg.get("val", "single"))
        edge.set(qn("w:sz"), str(cfg.get("sz", 6)))
        edge.set(qn("w:space"), "0")
        edge.set(qn("w:color"), cfg.get("color", "D5DCE4"))


def set_table_geometry(table, widths_dxa, indent=120):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent))
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        gc = OxmlElement("w:gridCol")
        gc.set(qn("w:w"), str(width))
        grid.append(gc)
    for row in table.rows:
        for cell, width in zip(row.cells, widths_dxa):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            cell.width = Inches(width / 1440)


def mark_header_row(row):
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def set_paragraph_shading(paragraph, fill):
    p_pr = paragraph._p.get_or_add_pPr()
    shd = p_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        p_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_paragraph_border(paragraph, edge="bottom", color="D5DCE4", size=8, space=5):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    b = OxmlElement(f"w:{edge}")
    b.set(qn("w:val"), "single")
    b.set(qn("w:sz"), str(size))
    b.set(qn("w:space"), str(space))
    b.set(qn("w:color"), color)
    p_bdr.append(b)


def set_repeat_keep(paragraph, keep_next=False, keep_lines=True):
    paragraph.paragraph_format.keep_together = keep_lines
    paragraph.paragraph_format.keep_with_next = keep_next


def add_field(paragraph, instruction):
    run = paragraph.add_run()
    fld_char = OxmlElement("w:fldChar")
    fld_char.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    sep = OxmlElement("w:fldChar")
    sep.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char, instr, sep, text, end])


def set_run_font(run, name="Calibri", size=None, color=None, bold=None, italic=None):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = rgb(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_image_alt(inline_shape, title, descr):
    doc_pr = inline_shape._inline.docPr
    doc_pr.set("title", title)
    doc_pr.set("descr", descr)


def setup_styles(doc):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = rgb(COLORS["body"])
    pf = normal.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(6)
    pf.line_spacing = 1.25

    specs = {
        "Title": (30, COLORS["night"], 0, 8),
        "Subtitle": (14, COLORS["muted"], 0, 10),
        "Heading 1": (16, COLORS["heart"], 18, 10),
        "Heading 2": (13, COLORS["indigo"], 14, 7),
        "Heading 3": (12, COLORS["indigo"], 10, 5),
        "Caption": (9, COLORS["muted"], 4, 8),
    }
    for name, (size, color, before, after) in specs.items():
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = rgb(color)
        style.font.bold = name in ("Title", "Heading 1", "Heading 2", "Heading 3")
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = name.startswith("Heading")
    styles["Caption"].font.italic = True


def add_numbering(doc):
    numbering = doc.part.numbering_part.element
    existing_abs = [int(e.get(qn("w:abstractNumId"))) for e in numbering.findall(qn("w:abstractNum"))]
    existing_num = [int(e.get(qn("w:numId"))) for e in numbering.findall(qn("w:num"))]
    abs_id = max(existing_abs or [0]) + 1
    num_id = max(existing_num or [0]) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abs_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "hybridMultilevel")
    abstract.append(multi)
    for ilvl, (fmt, text, left, hanging) in enumerate([
        ("bullet", "•", 540, 270), ("bullet", "-", 900, 270),
    ]):
        lvl = OxmlElement("w:lvl")
        lvl.set(qn("w:ilvl"), str(ilvl))
        start = OxmlElement("w:start"); start.set(qn("w:val"), "1"); lvl.append(start)
        num_fmt = OxmlElement("w:numFmt"); num_fmt.set(qn("w:val"), fmt); lvl.append(num_fmt)
        lvl_text = OxmlElement("w:lvlText"); lvl_text.set(qn("w:val"), text); lvl.append(lvl_text)
        jc = OxmlElement("w:lvlJc"); jc.set(qn("w:val"), "left"); lvl.append(jc)
        p_pr = OxmlElement("w:pPr")
        tabs = OxmlElement("w:tabs"); tab = OxmlElement("w:tab"); tab.set(qn("w:val"), "num"); tab.set(qn("w:pos"), str(left)); tabs.append(tab); p_pr.append(tabs)
        ind = OxmlElement("w:ind"); ind.set(qn("w:left"), str(left)); ind.set(qn("w:hanging"), str(hanging)); p_pr.append(ind)
        spacing = OxmlElement("w:spacing"); spacing.set(qn("w:after"), "80"); spacing.set(qn("w:line"), "300"); spacing.set(qn("w:lineRule"), "auto"); p_pr.append(spacing)
        lvl.append(p_pr)
        r_pr = OxmlElement("w:rPr"); color = OxmlElement("w:color"); color.set(qn("w:val"), COLORS["heart"]); r_pr.append(color); lvl.append(r_pr)
        abstract.append(lvl)
    numbering.append(abstract)
    num = OxmlElement("w:num"); num.set(qn("w:numId"), str(num_id)); abs_ref = OxmlElement("w:abstractNumId"); abs_ref.set(qn("w:val"), str(abs_id)); num.append(abs_ref); numbering.append(num)
    return num_id


def configure_page(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    section.different_first_page_header_footer = True

    header = section.header
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    hp.paragraph_format.space_after = Pt(3)
    r = hp.add_run("DEAD WALLS  |  GAME DESIGN BLUEPRINT  |  OWNER-APPROVED")
    set_run_font(r, size=8.5, color=COLORS["muted"], bold=True)
    set_paragraph_border(hp, "bottom", COLORS["stone"], 6, 4)

    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    rr = fp.add_run("DEAD WALLS  |  v1.0  |  PAGE ")
    set_run_font(rr, size=8.5, color=COLORS["muted"], bold=True)
    add_field(fp, "PAGE")


def add_heading(doc, text, level=1, kicker=None):
    if kicker:
        kp = doc.add_paragraph()
        kp.paragraph_format.space_before = Pt(0)
        kp.paragraph_format.space_after = Pt(3)
        kr = kp.add_run(kicker.upper())
        set_run_font(kr, size=8.5, color=COLORS["fire"], bold=True)
        set_repeat_keep(kp, True)
    p = doc.add_paragraph(text, style=f"Heading {level}")
    if level == 1:
        set_paragraph_border(p, "bottom", COLORS["stone"], 6, 5)
    return p


def add_para(doc, text, bold_lead=None, italic=False, color=None, align=None, after=None):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    if after is not None:
        p.paragraph_format.space_after = Pt(after)
    if bold_lead and text.startswith(bold_lead):
        r1 = p.add_run(bold_lead)
        set_run_font(r1, bold=True, color=color or COLORS["body"])
        r2 = p.add_run(text[len(bold_lead):])
        set_run_font(r2, italic=italic, color=color or COLORS["body"])
    else:
        r = p.add_run(text)
        set_run_font(r, italic=italic, color=color or COLORS["body"])
    return p


def add_bullet(doc, text, num_id, level=0, color=None):
    p = doc.add_paragraph()
    p_pr = p._p.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl"); ilvl.set(qn("w:val"), str(level))
    numid = OxmlElement("w:numId"); numid.set(qn("w:val"), str(num_id))
    num_pr.extend([ilvl, numid]); p_pr.append(num_pr)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.25
    r = p.add_run(text)
    set_run_font(r, color=color or COLORS["body"])
    return p


def add_callout(doc, label, text, accent="heart", fill="pale_red"):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [9360], indent=190)
    cell = table.cell(0, 0)
    set_cell_shading(cell, COLORS[fill])
    set_cell_margins(cell, top=150, start=190, bottom=150, end=190)
    set_cell_border(cell, left={"color": COLORS[accent], "sz": 22}, top={"color": COLORS[fill], "sz": 1}, bottom={"color": COLORS[fill], "sz": 1}, right={"color": COLORS[fill], "sz": 1})
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    lr = p.add_run(label.upper() + "  ")
    set_run_font(lr, size=9, color=COLORS[accent], bold=True)
    tr = p.add_run(text)
    set_run_font(tr, size=10.5, color=COLORS["ink"], bold=False)
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(1)
    spacer.paragraph_format.space_before = Pt(0)
    return table


def add_table(doc, headers, rows, widths, header_fill="indigo", compact=False):
    table = doc.add_table(rows=1, cols=len(headers))
    set_table_geometry(table, widths, indent=130)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    hdr = table.rows[0]
    mark_header_row(hdr)
    for i, text in enumerate(headers):
        cell = hdr.cells[i]
        set_cell_shading(cell, COLORS[header_fill])
        set_cell_margins(cell, 100 if compact else 120, 130, 100 if compact else 120, 130)
        set_cell_border(cell, top={"color":"FFFFFF","sz":1}, bottom={"color":"FFFFFF","sz":1}, left={"color":"FFFFFF","sz":1}, right={"color":"FFFFFF","sz":1})
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER if len(text) < 18 else WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(text)
        set_run_font(r, size=9.2 if compact else 9.6, color=COLORS["white"], bold=True)
    for row_idx, row in enumerate(rows):
        cells = table.add_row().cells
        fill = "FFFFFF" if row_idx % 2 == 0 else COLORS["paper"]
        for i, val in enumerate(row):
            cell = cells[i]
            set_cell_shading(cell, fill)
            set_cell_margins(cell, 95 if compact else 120, 130, 95 if compact else 120, 130)
            set_cell_border(cell, top={"color":COLORS["stone"],"sz":4}, bottom={"color":COLORS["stone"],"sz":4}, left={"color":COLORS["stone"],"sz":4}, right={"color":COLORS["stone"],"sz":4})
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.08
            if i > 0 and len(str(val)) < 22:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(str(val))
            set_run_font(r, size=9.2 if compact else 9.7, color=COLORS["body"], bold=(i == 0))
    doc.add_paragraph().paragraph_format.space_after = Pt(1)
    return table


def add_figure(doc, path: Path, caption: str, alt: str, width=6.5):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(3)
    shape = p.add_run().add_picture(str(path), width=Inches(width))
    set_image_alt(shape, caption, alt)
    cp = doc.add_paragraph(caption, style="Caption")
    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cp.paragraph_format.keep_with_next = False


def page_break(doc):
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)


def add_status_strip(doc, items):
    table = doc.add_table(rows=1, cols=len(items))
    widths = [9360 // len(items)] * len(items)
    widths[-1] += 9360 - sum(widths)
    set_table_geometry(table, widths)
    palette = [("pale_red","heart"),("pale_blue","frost"),("pale_gold","fire"),("pale_green","green")]
    for i,(value,label) in enumerate(items):
        cell=table.cell(0,i); fill,accent=palette[i%len(palette)]
        set_cell_shading(cell,COLORS[fill]); set_cell_margins(cell,120,120,120,120)
        set_cell_border(cell,top={"color":COLORS[accent],"sz":10},bottom={"color":COLORS["white"],"sz":1},left={"color":COLORS["white"],"sz":1},right={"color":COLORS["white"],"sz":1})
        cell.vertical_alignment=WD_ALIGN_VERTICAL.CENTER
        p=cell.paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_after=Pt(2)
        r=p.add_run(value); set_run_font(r,size=15,color=COLORS[accent],bold=True)
        p2=cell.add_paragraph(); p2.alignment=WD_ALIGN_PARAGRAPH.CENTER; p2.paragraph_format.space_after=Pt(0)
        r2=p2.add_run(label); set_run_font(r2,size=8.5,color=COLORS["muted"],bold=True)
    doc.add_paragraph().paragraph_format.space_after=Pt(1)


def add_source_link(doc, label, url):
    p=doc.add_paragraph()
    r=p.add_run(f"{label}: {url}")
    set_run_font(r,size=9.5,color=COLORS["indigo"])
    return p


def begin_page(doc, title, kicker, lead=None):
    add_heading(doc, title, 1, kicker)
    if lead:
        p = add_para(doc, lead, color=COLORS["ink"])
        p.paragraph_format.space_after = Pt(10)
    return p if lead else None


def build_document(fig):
    doc = Document()
    setup_styles(doc)
    configure_page(doc)
    bullet_num = add_numbering(doc)

    props = doc.core_properties
    props.title = "Dead Walls - Game Design Blueprint v1.0"
    props.subject = "Owner-approved game design and implementation blueprint"
    props.author = "Dead Walls Design"
    props.keywords = "Dead Walls, Unity, DOTS, ECS, incremental, roguelike, castle defense"
    props.comments = "Created as a clean document. Previous milestone document was not modified or reused."

    # COVER
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(5)
    r = p.add_run("OWNER-APPROVED GAME DESIGN BLUEPRINT")
    set_run_font(r, name="Arial", size=9.5, color=COLORS["fire"], bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run("DEAD WALLS")
    set_run_font(r, name="Arial", size=31, color=COLORS["night"], bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(10)
    r = p.add_run("Infinite Siege | Incremental Economy | Council Decisions | Procedural Castle Heart")
    set_run_font(r, name="Arial", size=13.5, color=COLORS["indigo"], bold=True)

    ip = doc.add_paragraph()
    ip.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ip.paragraph_format.space_after = Pt(8)
    shape = ip.add_run().add_picture(str(COVER), width=Inches(6.5))
    set_image_alt(shape, "Dead Walls cover art", "Fixed castle and archers face a massive single-type zombie horde under a day-night sky.")

    quote = doc.add_paragraph()
    quote.alignment = WD_ALIGN_PARAGRAPH.CENTER
    quote.paragraph_format.space_before = Pt(2)
    quote.paragraph_format.space_after = Pt(8)
    rr = quote.add_run("A fixed wall. One enemy. Ten thousand bodies. One run.")
    set_run_font(rr, name="Arial", size=13, color=COLORS["heart"], bold=True, italic=True)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.paragraph_format.space_after = Pt(4)
    rr = meta.add_run("PC / STEAM PRODUCT FRAME  |  UNITY 6 DOTS/ECS  |  VERSION 1.0  |  12 JULY 2026")
    set_run_font(rr, name="Arial", size=8.5, color=COLORS["muted"], bold=True)
    note = doc.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    nr = note.add_run("Sıfırdan hazırlanmıştır. Önceki milestone dokümanı güncellenmemiş ve içerik kaynağı olarak kullanılmamıştır.")
    set_run_font(nr, size=8.5, color=COLORS["muted"], italic=True)
    page_break(doc)

    # 01 AUTHORITY
    begin_page(doc, "Belge Otoritesi ve Okuma Kılavuzu", "Front Matter 01", "Bu belge, repo gerçeği ile owner tarafından tek tek onaylanan yeni tasarım kararlarını bir araya getiren geliştirme otoritesidir. Takvim değildir; ürün sözleşmesi, sistem kuralları ve kabul kapıları sunar.")
    add_callout(doc, "Belge kuralı", "Eski DEAD_WALLS_NEW_FEATURE_MILESTONE_PLAN dosyaları olduğu gibi kalır. Bu blueprint ayrı ad, ayrı görsel sistem ve ayrı içerik omurgasıyla yaratılmıştır.", "heart", "pale_red")
    add_heading(doc, "Karar statüleri", 2)
    add_table(doc,
              ["Statü", "Anlamı", "Belgedeki kullanım"],
              [
                  ("LOCKED", "Owner tarafından onaylandı", "Uygulama bunu değiştirmeden taşır"),
                  ("DESIGN TARGET", "Ölçülebilir ürün hedefi", "Test sonucuna göre teknik yöntem değişebilir"),
                  ("TUNING", "Formül veya sayı henüz dengelenmedi", "Inspector / SO üzerinden ayarlanır"),
                  ("POLISH OPEN", "İşlev kilitli, sunum dili açık", "Mockup ve görsel iterasyon gerektirir"),
              ], [1450, 3000, 4910], compact=True)
    add_heading(doc, "Bu belgenin çözmediği şeyler", 2)
    for text in [
        "Yeni hikaye, karakter kadrosu veya isimlendirilmiş antagonist icat etmez.",
        "Boss, miniboss, elit, özel gece veya ikinci düşman prefabı önermez.",
        "Mobil reklam / IAP planı kurmaz; etkileşim dili PC/Steam odaklıdır.",
        "Eski Mobile* sınıf adlarını sırf ürün yönü değişti diye yeniden adlandırma işi çıkarmaz.",
        "Kesin üretim süresi vermez. Uygulama sırası bağımlılığa göre tanımlanır.",
    ]:
        add_bullet(doc, text, bullet_num)
    add_callout(doc, "Tasarım pusulası", "Zero Stress King: Idle Defense referansı sürekli otomatik saldırı ve incremental büyüme için dokunma noktasıdır. Dead Walls; gerçek ölüm, 60 saniyelik gece-gündüz ritmi, işçi ekonomisi ve procedural Castle Heart ile ayrışır.", "frost", "pale_blue")
    add_source_link(doc, "Owner reference", "https://store.steampowered.com/app/4271160/Zero_Stress_King_Idle_Defense/")
    page_break(doc)

    # 02 NAVIGATION
    begin_page(doc, "Blueprint Haritası", "Front Matter 02", "Doküman dört katmanda okunur: ürün sözleşmesi, oyuncu sistemleri, teknik sahiplik ve üretim kapıları.")
    add_status_strip(doc, [("01", "ÜRÜN"), ("02", "SİSTEMLER"), ("03", "TEKNİK"), ("04", "TESLİM KAPILARI")])
    add_table(doc,
              ["Katman", "Sorusu", "Çıktı"],
              [
                  ("PRODUCT", "Oyuncu ne yaşar?", "Kimlik, core loop, koşu ve savaş sözleşmesi"),
                  ("SYSTEMS", "Sistemler nasıl çalışır?", "Ekonomi, Council, okçular, Heart, yetenekler ve meta"),
                  ("PRESENTATION", "Ekran nasıl okunur?", "HUD, onboarding, görsel ve işitsel yön"),
                  ("TECHNICAL", "Kod nerede değişir?", "Data, save, ECS/Mono sınırı ve performans"),
                  ("DELIVERY", "Ne zaman tamam sayılır?", "Uygulama sırası, test, risk ve DoD"),
                  ("EVIDENCE", "Hangi kaynaklar incelendi?", "Read-only source audit"),
              ], [1400, 3000, 4960])
    add_heading(doc, "Doküman içi ana terimler", 2)
    add_table(doc,
              ["Terim", "Bu belgede anlamı"],
              [
                  ("Run / Koşu", "Duvar ayaktayken devam eden ve Game Over ile sıfırlanan aktif oyun"),
                  ("Castle Heart", "Run teknolojisinin procedural graph ekranı ve Grave Essence harcama yüzeyi"),
                  ("Council", "Her üç günde bir toplanan koşu içi yönetim ve risk kararları katmanı"),
                  ("Meta", "Ölümden sonra kalan, sonraki koşuya ivme veren kalıcı upgrade listesi"),
                  ("Grave Essence", "Yalnızca run içinde teknoloji için kullanılan ve ölümde silinen kaynak"),
                  ("Spawn budget", "Düşman prefabını güçlendirmek yerine sahaya çıkarılacak adet baskısını yöneten değer"),
              ], [1900, 7460])
    page_break(doc)

    # 03 EXECUTIVE
    begin_page(doc, "Executive Design Contract", "Section 01 | Product", "Dead Walls bir tower defense içerik paketi değil; sayının kendisini tehdit ve görsel vaat haline getiren, tek cepheli incremental survival ürünüdür.")
    add_status_strip(doc, [("1", "DÜŞMAN PREFABI"), ("10.000", "AKTİF DÜŞMAN HEDEFİ"), ("1.000", "ORTAK OKÇU CAP"), ("60 sn", "TAM DÖNGÜ")])
    add_heading(doc, "Tek cümlelik ürün vaadi", 2)
    add_callout(doc, "Pitch", "Sabit bir kaleyi, büyüyen işçi ekonomisini ve 1.000 kişilik okçu garnizonunu yönet; aynı zombi prefabının on binler halinde aktığı sonsuz gecelerde procedural Castle Heart build'inle önceki rekorunu aş.", "heart", "pale_red")
    add_heading(doc, "Dört pazarlık dışı sütun", 2)
    pillars = [
        ("Sayı tehdittir", "Zombi istatistiği şişmez; adet, akış ve ekran yoğunluğu büyür."),
        ("Yönetim savaştan önce gelir", "Okçular otomatik savaşır; oyuncu işçi, kaynak, duvar, ammo ve teknoloji yönetir."),
        ("Her koşu farklı büyür", "Temel erişimler garantidir; Heart node'ları ve bağlantıları seed ile değişir."),
        ("Ölüm gerçektir", "Wall HP sıfır olduğunda run biter; gönüllü reset, ikinci can barı veya geri yükleme yoktur."),
    ]
    add_table(doc, ["Sütun", "Tasarım sözleşmesi"], pillars, [2200, 7160])
    add_heading(doc, "Ayrışma", 2)
    for text in [
        "Referans oyundaki otomatik akış korunur; baskısız / kayıpsız yapı korunmaz.",
        "Gece-gündüz yalnızca kozmetik değildir; spawn yoğunluğu ve Wall onarım erişimini değiştirir.",
        "Roguelike öğesi düşman rotası seçmek değil, run graph'ının ve build yollarının değişmesidir.",
        "Oyuncu gelen düşmanı seçmez, scout etmez ve anlamsız ön bilgi toplamaz.",
    ]:
        add_bullet(doc, text, bullet_num)
    page_break(doc)

    # 04 IDENTITY
    begin_page(doc, "Oyuncu Rolü ve Deneyim Tezi", "Section 01 | Product", "Oyuncu bir nişancı veya saha komutanı değil; duvarın arkasındaki insan, üretim ve teknoloji sistemlerinin karar sahibidir.")
    add_heading(doc, "Oyuncunun fiilleri", 2)
    add_table(doc,
              ["Sıklık", "Oyuncu fiili", "Sistem karşılığı"],
              [
                  ("Sürekli", "Kaynakları ve ok stokunu izler", "Üst HUD + ammo satın alma"),
                  ("Sık", "İşçi oranlarını değiştirir", "Farm / Lumberyard / Quarry / Mine"),
                  ("Sık", "Okçu alır veya yeniden eğitir", "Basic / Rapid / Frost ortak 1.000 cap"),
                  ("Dönemsel", "Duvarı gündüz onarır", "Stone ile tek seferlik repair"),
                  ("Dönemsel", "Council kararını verir", "3/6/9... günlerinin Dawn başlangıcı"),
                  ("Dönemsel", "Castle Heart node'u alır", "Grave Essence + pause"),
                  ("Taktik", "Fireball, Rally veya Emergency Repair kullanır", "Alt orta cooldown barı"),
              ], [1400, 3300, 4660])
    add_heading(doc, "Duygusal ritim", 2)
    for label, text in [
        ("Gündüz", "Kontrol, hız ve yatırım. Kale içi canlı görünür; üretim kararları okunur."),
        ("Akşam", "Kısa gerilim köprüsü. Büyük UI değil, renk ve ses oyuncuyu hazırlar."),
        ("Gece", "Aynı siluetin korkutucu sayıya dönüşmesi. Duvar normal onarıma kapanır."),
        ("Şafak", "Nefes, nüfus gelişi ve daha yüksek yeni taban. 3/6/9... günlerinde Council toplanır; zafer değil, bir sonraki baskının başlangıcıdır."),
    ]:
        add_callout(doc, label, text, "fire" if label in ("Gündüz","Şafak") else "frost", "pale_gold" if label in ("Gündüz","Şafak") else "pale_blue")
    add_heading(doc, "Narrative guardrail", 2)
    add_para(doc, "Castle Heart, Grave Essence ve sonsuz kuşatma mekanik olarak anlamlıdır; fakat dünyanın kökeni, oyuncunun unvanı ve zombilerin kaynağı bu tasarım turunda kilitlenmemiştir. Uygulama, onaylanmamış lore'u sistem sözleşmesine gömmemelidir.")
    page_break(doc)

    # 05 CORE LOOP
    begin_page(doc, "60 Saniyelik Kesintisiz Kuşatma", "Section 02 | Core Loop", "Döngü bir wave ekranı değil, çevresel ritimdir. Düşman akışı hiçbir fazda sıfıra inmez.")
    add_figure(doc, fig["core_loop"], "Şekil 1. Onaylı 60 saniyelik faz dağılımı.", "Circular day, dusk, night and dawn timeline with approved durations.")
    add_table(doc,
              ["Faz", "Süre", "Spawn", "Oyuncu erişimi"],
              [
                  ("Day", "30 sn", "Düşük ama artan", "İşçi, satın alma, normal Wall repair"),
                  ("Dusk", "5 sn", "Hızlanan geçiş", "Aynı erişim, güçlü görsel/ses sinyali"),
                  ("Night", "20 sn", "En yüksek akış", "Normal repair kapalı, aktif yetenekler açık"),
                  ("Dawn", "5 sn", "Hızlı düşüş", "Yeni nüfus gelişi ve yeni gün tabanı"),
              ], [1500, 1100, 2200, 4560], compact=True)
    add_callout(doc, "Locked", "Oyun hızı daima 1x'tir. x2/x4, gece atlama, offline üretim veya menüde güvenli kaynak kasma yoktur.", "heart", "pale_red")
    page_break(doc)

    # 06 RUN FLOW
    begin_page(doc, "Sonsuz Koşu, Gerçek Ölüm", "Section 02 | Run", "Koşunun final dalgası yoktur. Hedef duvar düşene kadar büyümek ve önceki rekoru aşmaktır.")
    add_figure(doc, fig["run_flow"], "Şekil 2. Aktif koşu, ana menü ve ölüm sonrası meta akışı.", "Flow showing active run can freeze in menu and only wall death ends the run before meta progression.")
    add_heading(doc, "Kayıt sözleşmesi", 2)
    for text in [
        "Tek aktif koşu bulunur; manuel save slot veya eski checkpoint yükleme yoktur.",
        "Ana menüye dönüş ve uygulamadan çıkış otomatik kayıt yapar.",
        "Continue; aynı gün, faz, kaynak, Wall HP, okçu sayısı ve generated graph ile devam eder.",
        "Wall HP sıfırlandığı anda ölüm kaydı kesinleşir; force-close ile geri alınamaz.",
        "Koşu ölmeden meta ödülü verilmez ve gönüllü prestige/reset bulunmaz.",
    ]:
        add_bullet(doc, text, bullet_num)
    add_callout(doc, "Game Over", "Duvar çökmez, içeri girme aşaması başlamaz ve Castle Heart ikinci can barına dönüşmez. Wall HP = 0 olduğunda simülasyon durur ve sonuç ekranı açılır.", "heart", "pale_red")
    page_break(doc)

    # 07 BATTLEFIELD
    begin_page(doc, "Tek Cepheli Savaş Alanı", "Section 03 | Battlefield", "Kale solda sabit, kuşatma alanı sağdadır. Oyuncu cephe, lane veya gelen düşman seçmez.")
    add_figure(doc, fig["battlefield"], "Şekil 3. Kale içi, duvar ve kuşatma alanının sabit kompozisyonu.", "Three-zone diagram of castle interior, wall and right-side siege field.")
    add_heading(doc, "Sahne kuralları", 2)
    for text in [
        "Düşmanlar sağ sınırdaki çoklu görünmez noktalardan çıkar ve duvar boyunca yayılır.",
        "Duvar görsel olarak geniştir; mekanik olarak tek ortak HP kullanır.",
        "Kale içindeki binalar önceden yerleştirilmiştir; build grid, drag, rotate veya placement yoktur.",
        "Dünya kamerası tamamen sabittir: pan, zoom ve rotation yoktur.",
        "Ekran oranı değişiminde kompozisyon korunur; kritik bölgeler kırpılmaz.",
    ]:
        add_bullet(doc, text, bullet_num)
    page_break(doc)

    # 08 HORDE
    begin_page(doc, "Tek Prefab Horde Tasarımı", "Section 03 | Horde", "Çıkış sürümü bilinçli olarak tek düşman prefabıyla çıkar. Monotonluğu boss veya varyantla değil, sayı koreografisi ve görsel okunurlukla çözer.")
    add_status_strip(doc, [("1", "PREFAB"), ("0", "BOSS"), ("0", "ELİT / VARYANT"), ("10.000+", "AKTİF HEDEF")])
    add_heading(doc, "Zorluk sözleşmesi", 2)
    for text in [
        "Gün ilerledikçe düşman HP, hasar ve hızına otomatik growth uygulanmaz.",
        "Zorluk, BaseSpawn(day) ile faz çarpanının oluşturduğu adet ve akış baskısından gelir.",
        "Aktif teknik limit dolarsa spawn talepleri silinmez; backlog içinde bekler ve boşluk oluşunca sahaya girer.",
        "Şafak yoğunluğu düşürür fakat koşuyu eski gün tabanına döndürmez.",
        "Prefab havuzu küçük başlayabilir ve ihtiyaçla genişleyebilir; ölen entity havuza döner.",
    ]:
        add_bullet(doc, text, bullet_num)
    add_figure(doc, fig["spawn_curve"], "Şekil 4. Temsili gün ve gece yoğunluk eğrisi; gerçek sayılar tuning değişkenidir.", "Log-scale illustrative curve where enemy count intensity grows while enemy stats stay fixed.")
    add_callout(doc, "Data-driven", "Çıkış katalogunda tek kayıt bulunur. Yeni düşman eklemek ileride prefab + tanım asset'i + katalog kaydı olmalı; ana spawn ve UI kodunda tür özel dal oluşmamalıdır.", "frost", "pale_blue")
    page_break(doc)

    # 09 WALL
    begin_page(doc, "Wall HP ve Onarım Ekonomisi", "Section 03 | Defense", "Duvar hem tek fail state hem de Stone için okunur, tek seferlik harcama yüzeyidir.")
    add_table(doc,
              ["Durum", "Kural", "UI"],
              [
                  ("Day / Dusk", "Stone ile normal onarım açık", "Minimal Wall bar + repair drawer"),
                  ("Night", "Normal onarım kapalı", "Emergency Repair cooldown dışında repair yok"),
                  ("HP = 0", "Anında Game Over", "Sonuç ekranı; çöküş / içeri girme yok"),
                  ("Max HP", "Castle Heart node'larıyla büyür", "Ayrı Gate/Core barı yok"),
              ], [1750, 4300, 3310])
    add_heading(doc, "Onarım formülü için tuning alanları", 2)
    for text in [
        "Repair paket büyüklüğü: sabit HP, yüzde HP veya hibrit formül.",
        "Eksik HP başına Stone fiyatı ve ilerleyen günlerde uygulanacak fiyat çarpanı.",
        "Emergency Repair yüzde değeri ve cooldown tabanı.",
        "Night başlangıcında açık normal repair işleminin iptal davranışı.",
        "Game Over ile aynı frame'de gelen repair girdisinin deterministik önceliği.",
    ]:
        add_bullet(doc, text, bullet_num)
    add_callout(doc, "Edge case", "Wall HP sıfıra değdiği frame Game Over kazanır. Aynı frame'de satın alınan repair veya cooldown tamamlanması ölümü geri çeviremez.", "heart", "pale_red")
    page_break(doc)

    # 10 ECONOMY
    begin_page(doc, "Incremental Ekonomi Sözleşmesi", "Section 04 | Economy", "Ana kaynaklar üretimden gelir ve yalnızca oyuncu bir satın alma yaptığında azalır. Bakım vergisi ve negatif üretim yoktur.")
    add_figure(doc, fig["economy"], "Şekil 5. Üretimden tek seferlik harcamalara kaynak akışı.", "Economy flow from four worker buildings to one-time purchases with no passive upkeep.")
    add_table(doc,
              ["Değer", "Kimliği", "Ana harcama alanları"],
              [
                  ("Food", "Nüfus büyümesi", "Şafakta gelen kişi, okçu alımı"),
                  ("Wood", "Hızlı büyüme ve ammo", "Yatak / kapasite, anında ok satın alma"),
                  ("Stone", "Savunma", "Gündüz Wall repair, savunma yatırımı"),
                  ("Iron", "İleri ekonomi", "Rapid/Frost ve ileri satın almalar"),
                  ("Arrows", "Tüketilen savaş stoğu", "Her atışta 1; Wood ile anında refill"),
                  ("Grave Essence", "Run teknolojisi", "Yalnızca Castle Heart node'ları"),
              ], [1500, 3100, 4760], compact=True)
    page_break(doc)

    # 11 WORKERS
    begin_page(doc, "Hazır Binalar ve İş Gücü", "Section 04 | Workers", "Oyuncu bina kurmaz. Dört hazır üretim binasına sayısal iş gücü atar; dünya bu kararın görsel yoğunluğunu temsil eder.")
    add_table(doc,
              ["Bina", "Üretim", "İşçi?", "Run yatırımı"],
              [
                  ("Farm", "Food", "Evet", "Kapasite + işçi verimliliği"),
                  ("Lumberyard", "Wood", "Evet", "Kapasite + işçi verimliliği"),
                  ("Quarry", "Stone", "Evet", "Kapasite + işçi verimliliği"),
                  ("Mine", "Iron", "Evet", "Kapasite + işçi verimliliği"),
                  ("Houses", "Yatak", "Hayır", "Tekrar satın alınan kapasite"),
                  ("Barracks", "Okçu satın alma", "Hayır", "Anında Basic/Rapid/Frost"),
                  ("Arcane Tower", "Büyü görseli", "Hayır", "Fireball dünyasal karşılık"),
              ], [1800, 2200, 1200, 4160], compact=True)
    add_heading(doc, "Dağıtım modeli", 2)
    for text in [
        "Oyuncu Food / Wood / Stone / Iron için kalıcı yüzde hedefleri belirler.",
        "Yeni nüfus şafakta hedef oranlara otomatik dağılır.",
        "Bina kapasitesi dolarsa fazlalık Idle Population havuzunda kalır.",
        "Oyuncu +1, +10, +100 ve doğrudan sayı düzeltmesi yapabilir.",
        "Atama ücretsiz, anlık ve oyun akarken yapılır.",
    ]:
        add_bullet(doc, text, bullet_num)
    add_heading(doc, "Dünya temsili", 2)
    add_para(doc, "İşçiler dünyada görünür; ancak sınırsız büyüyen sayısal nüfus bire bir render edilmez. Görünen worker entity yoğunluğu, bina animasyonları, fenerler, taşıma ve üretim feedback'i gerçek tahsisi temsil eder. Gameplay truth sayısal allocation state'tir.")
    add_callout(doc, "Removed", "Fletcher gameplay binası ve Fletcher işçisi yoktur. Blacksmith / Wizard Academy sırf sahnede var diye yeni sistem üretmez; onaylanana kadar dekor veya dormant content kalabilir.", "heart", "pale_red")
    page_break(doc)

    # 12 POPULATION
    begin_page(doc, "Nüfus, Yatak ve Şafak Gelişi", "Section 04 | Population", "Nüfus incremental büyür; bakım tüketimiyle geri alınmaz. Büyümenin kapısı yatak ve tek seferlik Food maliyetidir.")
    add_heading(doc, "Onaylı kurallar", 2)
    for text in [
        "Yatak kapasitesinin sabit üst sınırı yoktur; maliyet sahip olunan kapasiteyle büyür.",
        "Yeni kurtulanlar şafakta sağ taraftan yürüyerek kaleye gelir ve nüfusa eklenir.",
        "Her gelen kişi Food'u yalnızca bir kez azaltır; mevcut nüfus pasif Food tüketmez.",
        "Food yetersizse gelebilecek kişi sayısı mevcut bütçeyle sınırlanır.",
        "Yatak yetersizse fazla nüfus kabul edilmez; kapasite hiçbir zaman aşılmaz.",
        "Açlık, nüfus ölümü, göç veya üretim cezası yoktur.",
    ]:
        add_bullet(doc, text, bullet_num)
    add_heading(doc, "Incremental fiyat davranışı", 2)
    add_para(doc, "Yatak, okçu, bina kapasitesi ve verimlilik satın almaları; sahip olunan adet veya seviyeye göre büyüyen maliyet eğrisi kullanır. Fiyatın 100'den 100.000'e çıkması hedef davranıştır. Kesin growth katsayısı veri odaklı tuning değişkenidir.")
    add_table(doc,
              ["Satın alma", "Fiyatı ne büyütür?", "Üst sınır"],
              [
                  ("Yatak", "Sahip olunan yatak", "Yok"),
                  ("Okçu", "Aynı türün mevcut sayısı", "Üç tür toplam 1.000"),
                  ("Bina kapasitesi", "Bina bazlı kapasite seviyesi", "Tasarım cap yok; sayı int güvenliğiyle korunur"),
                  ("İşçi verimliliği", "Bina bazlı verim seviyesi", "Tasarım cap yok; formül soft-cap kullanabilir"),
              ], [1900, 3900, 3560])
    page_break(doc)

    # 13 COUNCIL
    begin_page(doc, "Council: Koşu İçi Karar Meclisi", "Section 04 | Council", "Council kaldırılan veya dormant bir yan özellik değildir. Castle Heart teknoloji build'ini kurarken Council, koşunun yaşayan ekonomik ve savunma kararlarını yönetir.")
    add_status_strip(doc, [("3 GÜN", "DÜZENLİ TAKVİM"), ("BAĞLAM", "DUYARLI SEÇİM"), ("2", "NET SEÇENEK"), ("CURATED", "ONAYLI İÇERİK")])
    add_table(doc,
              ["Toplantı", "Tetik", "Takvim davranışı", "Karar rolü"],
              [
                  ("Düzenli Council", "3, 6, 9, 12... günlerinin Dawn başlangıcı", "Kesindir; aynı gün yalnız bir kez açılır", "Ekonomi, nüfus, Wall ve sonraki gece arasında anlamlı trade-off"),
              ], [1700, 2450, 2600, 2610], compact=True)
    add_heading(doc, "İçerik üretim sözleşmesi", 2)
    for text in [
        "Olaylar serbest biçimli runtime yapay zekâ metniyle üretilmez. İnsan tarafından yazılmış ve onaylanmış şablonlar ile etkiler, mevcut CouncilComposer tarafından deterministik biçimde bağlama göre birleştirilir.",
        "Her seçenek seçilmeden önce gerçek sayısal etkisini açıkça gösterir; belirsiz veya yalnızca kozmetik seçim kartı kullanılmaz.",
        "Kaynak kıtlığı, üretim hızı, Wall durumu ve önceki Council kararları bağlam seçiminde kullanılabilir; aynı şablonun anlamsız tekrarı engellenir.",
        "Olay zincirleri yalnız editoryal incelemeden geçmiş flag bağlantılarıyla açılır. Yeni lore, karakter veya fraksiyon owner onayı olmadan canon yapılmaz.",
    ]:
        add_bullet(doc, text, bullet_num)
    add_heading(doc, "Ana sistem guardrail'leri", 2)
    for text in [
        "Nüfus kazancı yatak kapasitesi ve tek seferlik Food sözleşmesini bypass edemez.",
        "Okçu kazancı toplam nüfus ve ortak 1.000 okçu cap'ini bypass edemez.",
        "Savunma etkisi yalnız Wall HP / Wall Max HP üzerinde çalışır; Gate veya Core geri getirmez.",
        "Gece baskısı etkisi düşman HP, hasar veya hızını değil yalnız adet ve akış çarpanını değiştirir.",
        "Düzenli takvim, seçilmiş seçenekler, zincir flag'leri ve aktif süreli etkiler exact run save içinde korunur.",
    ]:
        add_bullet(doc, text, bullet_num)
    add_callout(doc, "Ownership", "Castle Heart = teknoloji build'i. Council = koşu içi yönetim ve risk kararları. Meta = ölüm sonrası kalıcı ilerleme. Bu üç yüzey birbirinin para birimini veya upgrade rolünü devralmaz.", "frost", "pale_blue")
    page_break(doc)

    # 14 ARCHERS
    begin_page(doc, "Üç Okçu, Tek Garnizon", "Section 05 | Archers", "Basic, Rapid ve Frost korunur; sert counter sistemi yoktur. Oyuncu ortak 1.000 kişilik kapasiteyi build tercihine göre böler.")
    add_status_strip(doc, [("BASIC", "DENGELİ"), ("RAPID", "HIZLI ATIŞ"), ("FROST", "YAVAŞLATMA"), ("1.000", "ORTAK CAP")])
    add_table(doc,
              ["Tür", "Run erişimi", "Doğal rol", "Özel kural"],
              [
                  ("Basic", "Koşu başında açık", "Dengeli hasar / menzil / hız", "Rapid veya Frost'a retrain edilebilir"),
                  ("Rapid", "Heart içinde garanti unlock", "Daha düşük vuruş, daha yüksek fire rate", "Daha hızlı arrow consumption"),
                  ("Frost", "Heart içinde garanti unlock", "Hasar + slow", "Slow gücü/duration Heart'tan büyür"),
              ], [1400, 2400, 3000, 2560])
    add_heading(doc, "Satın alma ve retraining", 2)
    for text in [
        "Okçu satın alımı anlıktır; eğitim kuyruğu veya Barracks worker yoktur.",
        "Her okçu toplam nüfustan kişi kullanır ve toplam ortak 1.000 cap'e sayılır.",
        "Rapid/Frost açıldığında mevcut Basic okçular tek seferlik kaynak maliyetiyle dönüştürülebilir.",
        "Tür maliyeti kendi mevcut sayısına göre artar; kesin formül tuning verisidir.",
        "Okçular ayrı HP kullanmaz, düşman tarafından öldürülmez ve Wall düşene kadar korunur.",
    ]:
        add_bullet(doc, text, bullet_num)
    add_callout(doc, "No duplicate upgrade UI", "Hasar, atış sıklığı, menzil ve türe özel bütün güçlendirmeler Castle Heart node'larında yaşar. Market / Barracks içinde ayrı archer level sistemi kurulmaz.", "frost", "pale_blue")
    page_break(doc)

    # 14 AMMO
    begin_page(doc, "Ok Stoğu: Tek Sürekli Tüketilen Savaş Değeri", "Section 05 | Ammo", "Ana kaynaklar pasif akmaz; ok stoğu ise her atışta azalır. Oyuncu Wood ödeyerek anında refill yapar.")
    add_table(doc,
              ["Eylem", "Kaynak davranışı", "Bekleme"],
              [
                  ("Ok atışı", "Arrow -1", "Archer fire rate"),
                  ("Ok satın alma", "Wood tek seferlik azalır, Arrow anında artar", "Yok"),
                  ("Arrow capacity upgrade", "Tek seferlik büyüyen maliyet", "Yok"),
                  ("Arrow efficiency upgrade", "Aynı Wood için daha fazla Arrow", "Castle Heart veya run purchase kararıyla bağlanır"),
              ], [2100, 4400, 2860])
    add_heading(doc, "Fiyat sözleşmesi", 2)
    for text in [
        "Okun birim fiyatı her refill işleminde sonsuza doğru büyümez.",
        "Daha büyük ordu ve daha yüksek fire rate doğal olarak daha fazla arrow talebi yaratır.",
        "Rapid'in gücü ok ekonomisinde görünür bedel üretir; düşman counter'ı gerekmez.",
        "Arrow 0 olduğunda okçular yeni stok gelene kadar ateş etmez.",
        "Refill input'u aynı frame'de stoğa yansır; üretim queue'su ve Fletcher yoktur.",
    ]:
        add_bullet(doc, text, bullet_num)
    add_callout(doc, "UI", "Ammo paneli tek satırda Current / Capacity, sabit oranlı paketler ve Alabildiğin Kadar satın alma kontrolü gösterir. Paket fiyatı satın alma öncesi açıkça görünür.", "fire", "pale_gold")
    page_break(doc)

    # 15 PLACEMENT
    begin_page(doc, "Okçu Yerleşimi: Mevcut Tilemap Fikrinin Ölçeklenmesi", "Section 05 | Placement", "Yeni bir 1.000 slot sistemi kurulmaz. Mevcut outside tilemap dizilimi 40 tile ve tile başına 25 kararlı noktayla tamamlanır.")
    add_figure(doc, fig["placement"], "Şekil 6. 40 tile ve tile başına 25 doğal, deterministik okçu noktası.", "Archer placement diagram showing forty tile overview and a zoomed diamond tile with twenty-five stable random points.")
    add_heading(doc, "Uygulama kuralları", 2)
    for text in [
        "Her tile için local 25 nokta, tile koordinatı + slot index seed'iyle üretilir.",
        "Örnekleme izometrik diamond içinde ve güvenli inset ile yapılır.",
        "Minimum local mesafe sprite ayaklarının üst üste binmesini azaltır.",
        "Doldurma layer mantığı: önce bütün 40 tile, sonra ikinci local slotlar.",
        "Save pozisyon saklamak zorunda değildir; count + stable algorithm aynı formasyonu yeniden üretir.",
        "Preview gizmo bütün 1.000 noktayı editor'da denetlenebilir göstermelidir.",
    ]:
        add_bullet(doc, text, bullet_num)
    page_break(doc)

    # 16 TARGETING
    begin_page(doc, "Otomatik Hedefleme ve Ölçek", "Section 05 | Combat", "Oyuncu okçu veya düşman seçmez. Her okçu kendisine en yakın geçerli düşmana ateş eder.")
    add_heading(doc, "Davranış sözleşmesi", 2)
    for text in [
        "Range içindeki yaşayan ve death state'e girmemiş en yakın düşman hedeflenir.",
        "Basic, Rapid ve Frost aynı hedefleme kuralını kullanır.",
        "Hedef seçimi Wall'a en yakın düşmana global olarak yığılmaz; okçu konumu doğal dağılım üretir.",
        "Mermi hedef ölür veya havuza dönerse deterministik şekilde temizlenir ya da yeni hedef alır; kesin politika tek olmalıdır.",
        "Aynı düşmana ayrılan tahmini incoming damage, gereksiz binlerce ok overkill'ini önler.",
    ]:
        add_bullet(doc, text, bullet_num)
    add_heading(doc, "Mevcut teknik risk", 2)
    add_callout(doc, "Scale blocker", "Current ArcherShootSystem her hazır okçu için bütün zombileri brute-force tarar. Kod yorumu yaklaşık 10 okçu × 6.000 zombi varsayar; yeni 1.000 × 10.000 sözleşmesi için spatial query, hedef rezervasyonu ve Burst/job ölçeği yeniden ele alınmalıdır.", "heart", "pale_red")
    add_table(doc,
              ["Kabul metriği", "Hedef"],
              [
                  ("Target search", "Archer başına bütün 10.000 düşmanı taramaz"),
                  ("Overkill", "Aynı frame'de ölüme yetecek damage üstü yığılma kontrollüdür"),
                  ("Determinism", "Save/load ve aynı seed kritik seçim sırasını bozmaz"),
                  ("Ammo truth", "Atılan gerçek projectile sayısı kadar Arrow azalır"),
              ], [2600, 6760])
    page_break(doc)

    # 17 HEART INTRO
    begin_page(doc, "Castle Heart: Run Teknolojisinin Tek Sahibi", "Section 06 | Castle Heart", "Castle Heart hem incremental stat büyümesini hem roguelike davranış değişimini taşıyan tek teknoloji yüzeyidir.")
    add_figure(doc, fig["tech"], "Şekil 7. Bir koşunun örnek Heart görünümü; node zinciri sabit değildir.", "Procedural Castle Heart example with four stable directions and variable nodes.")
    add_heading(doc, "Ekran ve etkileşim", 2)
    for text in [
        "HUD'daki Castle Heart butonu tam ekran graph açar ve bütün simülasyonu durdurur.",
        "Mouse drag ile pan ve wheel ile graph zoom yalnızca bu ekranda çalışır.",
        "Node satın alma para birimi yalnızca Grave Essence'tır.",
        "İlk satın alma node'u Level 1 yapar ve bağlı komşuları reveal eder.",
        "Tekrarlanabilir node +1 / +10 / Alabildiğin Kadar kontrolleri sunar.",
    ]:
        add_bullet(doc, text, bullet_num)
    add_callout(doc, "Core guarantee", "Basic run başında açıktır. Rapid, Frost ve Fireball her generated graph içinde erişilebilir olmak zorundadır; rastgelelik temel sistemi kaybettiremez.", "frost", "pale_blue")
    page_break(doc)

    # 18 HEART GENERATION
    begin_page(doc, "Graph Üretim Sözleşmesi", "Section 06 | Procedural Graph", "Ağ reveal anında zar atmaz. Run başladığında tamamen üretilir, gizli tutulur ve kesin graph save'e yazılır.")
    add_heading(doc, "Sabit olan", 2)
    for text in [
        "Castle Heart merkez node'u.",
        "Ordu, Savunma, Üretim ve Heart/Büyü yönlerinin görsel pusulası.",
        "Rapid, Frost, Fireball ve temel Wall erişiminin varlığı.",
        "Her ana yönde en az bir tekrar geliştirilebilir harcama noktası.",
    ]:
        add_bullet(doc, text, bullet_num)
    add_heading(doc, "Koşuya göre değişen", 2)
    for text in [
        "Node listesi, bağlantılar ve erişim sırası.",
        "Nadir evolution node'larının bulunması.",
        "Keystone çiftleri ve graph derinliği.",
        "Aynı temel unlock'a kısa pahalı veya uzun zengin rotayla ulaşılması.",
    ]:
        add_bullet(doc, text, bullet_num)
    add_heading(doc, "Üretim adımları", 2)
    add_table(doc,
              ["Adım", "İşlem", "Validation"],
              [
                  ("1", "Run seed ve dört yön iskeleti", "Root tek ve bütün yönler bağlı"),
                  ("2", "Garanti node'ları izinli derinliğe yerleştir", "Rapid/Frost/Fireball reachable"),
                  ("3", "Tag + rarity + depth ile havuzu doldur", "Duplicate / invalid prereq yok"),
                  ("4", "Edge, cross-link ve Keystone çiftlerini kur", "Normal node yanlışlıkla lock üretmez"),
                  ("5", "Graph'i doğrula ve save'e yaz", "Disconnected / dead core path yok"),
              ], [1000, 5000, 3360], compact=True)
    page_break(doc)

    # 19 NODE TYPES
    begin_page(doc, "Node Türleri, Maliyet ve Reveal", "Section 06 | Node Design", "Graph'ın procedural olması node'ların belirsiz veya anlamsız olmasına izin vermez. Oyuncu gördüğü node'un etkisini tam bilir.")
    add_table(doc,
              ["Node türü", "Satın alma", "Örnek", "Bağlantı etkisi"],
              [
                  ("Unlock", "Tek sefer", "Rapid, Frost, Fireball", "Sistemi ve devam node'larını açar"),
                  ("Repeatable", "Sınırsız / soft-cap", "Damage, fire rate, Wall HP", "İlk level reveal; sonrası stat büyütür"),
                  ("Evolution", "Tek sefer", "Split shot, burning ground", "Davranışı değiştirir"),
                  ("Keystone", "Tek sefer", "Karşıt doctrine", "Yalnız eş Keystone'u kapatır"),
              ], [1600, 2000, 2600, 3160], compact=True)
    add_heading(doc, "Reveal ilkesi", 2)
    for text in [
        "Başlangıçta Heart ve ona bağlı ilk seçenekler tamamen görünür.",
        "Uzakta yalnızca yönün rengi / damarı görünür; exact node gizlidir.",
        "Gizli içerik run başında zaten belirlenmiştir; save-scum ile reroll olmaz.",
        "Bir node ilk kez alındığında yalnız bağlı komşular açılır.",
        "Keystone görünür olduğunda karşıt seçimi ve kapanacak node açıkça işaretlenir.",
    ]:
        add_bullet(doc, text, bullet_num)
    add_heading(doc, "Stat sınırları", 2)
    add_para(doc, "Damage ve maliyet gibi değerler büyük sayılara çıkabilir. Fire rate, cooldown, slow ve range fiziksel/teknik limitlere yaklaştığında soft-cap veya diminishing return kullanır; UI node'u satın alınamaz hale getirmek yerine kalan değeri açıkça gösterir.")
    add_callout(doc, "No parallel system", "Mevcut Archer type level paneli veya ayrı Blacksmith tech yüzeyi run progression otoritesi değildir. Yeni numeric ve davranışsal okçu upgrade'leri Castle Heart effect pipeline'ında toplanır.", "heart", "pale_red")
    page_break(doc)

    # 20 ABILITIES
    begin_page(doc, "Aktif Yetenek Barı", "Section 07 | Active Play", "Oyuncu savaşın sahibidir ama mikro hedefleme yapmaz. Üç cooldown farklı problemi çözer ve alt orta HUD'da yaşar.")
    add_table(doc,
              ["Yetenek", "Input", "Etki", "Kaynak"],
              [
                  ("Fireball", "1 + dünya alanı seçimi", "Seçilen yarıçapta AoE damage", "Yok; cooldown"),
                  ("Rally", "2", "Bütün okçulara kısa fire rate boost", "Yok; cooldown"),
                  ("Emergency Repair", "3", "Night sırasında Wall HP yüzdesi geri verir", "Yok; uzun cooldown"),
              ], [1900, 2300, 3300, 1860])
    add_heading(doc, "Fireball ilk büyüdür", 2)
    for text in [
        "Fireball her run graph'ında garanti bulunur; koşu başında açık olmak zorunda değildir.",
        "Damage, radius ve cooldown repeatable Heart node'larıyla büyür.",
        "Burning ground veya second blast gibi evolution'lar havuzda nadir çıkabilir.",
        "Mana veya ana kaynak maliyeti yoktur; targeting sırasında UI click'leri cast sayılmaz.",
        "Yeni büyüler ileride meta pool unlock olarak eklenebilir; çıkışta yalnız Fireball garantidir.",
    ]:
        add_bullet(doc, text, bullet_num)
    add_callout(doc, "Removed", "Arrow Storm aktif barından çıkarılmıştır. Fireball zaten alan saldırısını, 1.000 okçu ise ok yağmuru fantezisini taşır.", "heart", "pale_red")
    page_break(doc)

    # 21 META
    begin_page(doc, "Ölüm Sonrası Meta İlerleme", "Section 08 | Meta", "Meta ikinci graph değildir. Game Over ekranında sade, kalıcı ve uzun vadeli harcama listesi sunar.")
    add_heading(doc, "Ödül hesabı", 2)
    add_para(doc, "Kazanılan meta para; ulaşılan gün, hayatta kalınan gece, toplam kill, en yüksek nüfus ve koşu rekorlarından türetilir. Boss ölçütü yoktur. Kesin ağırlıklar tuning verisidir ve tek run'da iki kez ödül yazılmamalıdır.")
    add_table(doc,
              ["Kalıcı upgrade", "Etki", "Heart ile sınır"],
              [
                  ("Starting resources", "Wood / Stone / Iron / Food", "Node açmaz"),
                  ("Starting Basic Archers", "İlk büyümeyi hızlandırır", "Rapid/Frost açmaz"),
                  ("Starting beds", "İlk şafak nüfusunu hızlandırır", "Run yatak fiyatını silmez"),
                  ("Base Wall HP", "Koşuya dayanıklılık tabanı", "Heart Wall node'ları yine değerlidir"),
                  ("Worker production", "Global küçük multiplier", "Run capacity/verim satın alımı sürer"),
                  ("Arrow efficiency", "Wood başına daha fazla Arrow", "Ammo kararı yok olmaz"),
                  ("Essence gain", "Heart progression hızlanır", "Graph sonucu değişmez"),
                  ("Node pool unlock", "Yeni olası evolution / spell", "Mevcut run'a zorla eklenmez"),
              ], [2400, 3400, 3560], compact=True)
    add_heading(doc, "Meta guardrail", 2)
    for text in [
        "Aktif koşu sürerken meta satın alınamaz.",
        "Meta, generated graph edge'lerini veya Keystone sonucunu seçmez.",
        "StartingTechLevel etkisi procedural graph'ı atladığı için yeni modelde kullanılmaz.",
        "Tekrarlanabilir meta sink'lerin maliyeti büyür; içerik unlock'ları tek seferliktir.",
    ]:
        add_bullet(doc, text, bullet_num)
    page_break(doc)

    # 22 PERSISTENCE
    begin_page(doc, "Run Save ve Meta Save Matrisi", "Section 08 | Persistence", "Run state ile kalıcı state iki ayrı otoritedir. Birbirine yanlış veri sızması ölüm sözleşmesini bozar.")
    add_table(doc,
              ["Veri", "Run save", "Meta save", "Game Over"],
              [
                  ("Gün / faz / cycle timer", "Evet", "Hayır", "Sil"),
                  ("Wood / Stone / Iron / Food / Arrow", "Evet", "Hayır", "Sil"),
                  ("Nüfus / worker oranı / yatak", "Evet", "Hayır", "Sil"),
                  ("Basic / Rapid / Frost count", "Evet", "Başlangıç bonusu ayrı", "Sil"),
                  ("Generated graph / node levels", "Evet", "Pool unlock ayrı", "Sil"),
                  ("Council takvimi / flags / aktif etkiler", "Evet", "Hayır", "Sil"),
                  ("Wall HP / ability cooldown", "Evet", "Base bonus ayrı", "Sil"),
                  ("Meta para / upgrade levels", "Hayır", "Evet", "Koru ve ödül ekle"),
                  ("Tutorial tamamlandı", "Hayır", "Evet", "Koru"),
              ], [2600, 1900, 2400, 2460], compact=True)
    add_heading(doc, "Kritik save edge case'leri", 2)
    for text in [
        "Main menu dönüşünde run graph'ın gizli node'ları dahil tam snapshot alınır.",
        "Save sırasında 10.000 düşmanın tek tek pozisyonu yerine deterministik yeniden kurulum stratejisi değerlendirilebilir; oyuncu durumunu değiştirmemelidir.",
        "Game Over işlemi meta ödülünü idempotent yazar; tekrar açılışta ikinci kez eklemez.",
        "Versiyon göçü eski Mobile save'lerini yeni run sözleşmesine sessizce yanlış map etmemelidir.",
    ]:
        add_bullet(doc, text, bullet_num)
    page_break(doc)

    # 23 HUD
    begin_page(doc, "Sabit Kamera ve Minimal HUD", "Section 09 | UX", "Dünya kompozisyonu sabit kalır. UI savaş alanını çerçeveler; üstüne oturup ana görseli bastırmaz.")
    add_figure(doc, fig["hud"], "Şekil 8. Onaylı HUD yerleşim bölgeleri; faz göstergesinin görsel dili henüz açık.", "HUD wireframe with resource strip, small phase placeholder, one wall bar, bottom abilities and management buttons.")
    add_heading(doc, "Kesin UI kararları", 2)
    for text in [
        "Main Camera: pan yok, zoom yok, rotation yok.",
        "Wall, Gate, Core üçlüsü kaldırılır; tek minimal Wall barı kalır.",
        "Fireball, Rally ve Emergency Repair ekranın alt ortasındadır.",
        "Workers / Housing alt sol; Archers / Castle Heart alt sağdır.",
        "Kaynaklar üst alanda kalır; aynı anda yalnız bir yönetim drawer'ı açık olabilir.",
        "Council kartı geçici karar yüzeyidir; kalıcı drawer olmaz, iki exact etkiyi ve karar süresini açıkça gösterir.",
        "Düşman forecast, horde pressure tahmini veya anlamsız pre-wave bilgi paneli yoktur.",
    ]:
        add_bullet(doc, text, bullet_num)
    add_callout(doc, "Polish open", "Mevcut büyük CyclePanel ve DAY / DUSK / NIGHT etiket dili reddedilmiştir. Üst ortada küçük bir alan ayrılır; kesin ikonografi, animasyon ve tipografi ayrı mockup turunda onaylanır.", "fire", "pale_gold")
    page_break(doc)

    # 24 ONBOARDING
    begin_page(doc, "İlk Koşu Onboarding", "Section 09 | First-Time Experience", "Tutorial yalnızca ilk koşuda, bağlam oluştuğunda ve kısa yönlendirmelerle çalışır. Sonraki koşuların ritmini kesmez.")
    add_table(doc,
              ["Tetik", "Öğretilen", "Sunum"],
              [
                  ("İlk gündüz", "İşçi oranı", "Workers düğmesi pulse + tek satır açıklama"),
                  ("İlk kaynak yeterliliği", "Basic Archer", "Archer drawer highlight"),
                  ("İlk düşük ammo", "Arrow satın alma", "Ammo satırı highlight; zorunlu popup yok"),
                  ("İlk kill / Essence", "Castle Heart", "Heart butonu pulse; graph açılınca oyun durur"),
                  ("İlk düzenli Council / Day 3", "Seçim, bedel ve sonuç", "İki etkinin exact sonucu gösterilir; seçim sonrası sonuç feedback'i"),
                  ("İlk Wall hasarı sonrası day", "Normal repair", "Wall bar + repair action highlight"),
                  ("İlk Night", "Aktif yetenek barı", "Unlock olan düğme üzerinde key hint"),
              ], [2450, 2300, 4610], compact=True)
    add_heading(doc, "Onboarding guardrail", 2)
    for text in [
        "Öğretim sistemi oyuncu adına kaynak harcamaz veya worker dağıtmaz.",
        "Oyunu sürekli durduran modal zinciri yoktur.",
        "Bir prompt görülmeden ilgili işlem yapılırsa adım tamamlanmış sayılır.",
        "Tutorial tamamlandı flag'i meta save'de kalır; Settings'ten reset edilebilir.",
        "Player-facing metin English olur; bu tasarım dokümanı Türkçedir.",
    ]:
        add_bullet(doc, text, bullet_num)
    page_break(doc)

    # 25 VISUAL AUDIO
    begin_page(doc, "Görsel ve İşitsel Creative Direction", "Section 10 | Polish", "Gece-gündüz oyunun ayırt edici yüzüdür. Faz UI'sı küçülürken dünya ışığı ve ses daha fazla sorumluluk alır.")
    add_table(doc,
              ["Katman", "Day", "Dusk", "Night", "Dawn"],
              [
                  ("Işık", "Sıcak, üretim okunur", "Hızlı amberden indigo'ya", "Soğuk ay, güçlü siluet", "Kısa cyan/altın kırılması"),
                  ("Kale", "Bina hareketi", "Fenerler yanar", "Pencereler ve ok salvoları", "Kapılar/nüfus gelişi"),
                  ("Horde", "Seyrek çizgi", "Yoğunlaşan edge", "Kütle / dalga dokusu", "Spawn düşer, kalanlar sürer"),
                  ("Ses", "İşçi ve ambiyans", "Gerilim riser", "Yoğun ama rate-limited mix", "Nefes + yeni gün cue"),
              ], [1300, 2015, 2015, 2015, 2015], compact=True)
    add_heading(doc, "10.000 düşman okunurluğu", 2)
    for text in [
        "Hit VFX ve SFX her düşmanda üretilemez; yoğunluğa göre budget / rate limit kullanır.",
        "Kalabalık tek koyu leke olmamalı; ground contrast, silhouette edge ve motion cadence korunur.",
        "Fireball ve Frost feedback'i horde içinde kaybolmayacak renk/ölçek hiyerarşisine sahip olur.",
        "Okçu salvoları tek tek projectile kaosu yerine okunur toplu ritim üretmelidir.",
        "Faz geçişleri tam ekran büyük yazı yerine color grading, sky, particles ve audio ile okunur.",
    ]:
        add_bullet(doc, text, bullet_num)
    add_callout(doc, "Not locked", "Kesin DAY/DUSK/NIGHT widget tasarımı, meta para adı, lore metni ve ek büyü listesi bu dokümanda bilinçli olarak final değildir.", "fire", "pale_gold")
    page_break(doc)

    # 26 ARCHITECTURE
    begin_page(doc, "Teknik Sahiplik ve Data Sınırları", "Section 11 | Technical", "Yeni ürün yönü mevcut ECS/Mono köprüsünü kullanır; var olan owner'ları paralel sistemlerle çoğaltmaz.")
    add_figure(doc, fig["architecture"], "Şekil 9. Definition asset, run runtime, presentation ve persistence sınırları.", "Data-driven architecture showing assets, ECS runtime, UI presentation and persistence.")
    add_table(doc,
              ["Karar alanı", "Mevcut owner / temas", "Yeni sorumluluk"],
              [
                  ("Cycle + spawn", "ContinuousSiegeCycleSystem, WaveSpawnSystem", "60 sn sabit faz ve quantity-only scaling"),
                  ("Düşman data", "Zombie prefab / authoring", "Tek kayıtlı EnemyDefinition + expandable pool"),
                  ("Workers", "MobilePopulationEconomySystem, worker visual", "4 oran + capacity + visual density"),
                  ("Archers", "GameManager, ArcherShootSystem", "1.000 ortak cap, scalable targeting"),
                  ("Placement", "MobileCastleArcherTilePlacement", "40×25 stable local points"),
                  ("Heart", "TechNodeDefinitionSO, TechTreeUI", "Run graph generator + Essence currency"),
                  ("Council", "CouncilComposer, CouncilEventUI, CouncilEventCatalog", "3/6/9... schedule + guarded effects + context memory"),
                  ("Meta", "MetaProgression", "Death-only reward + fixed upgrade list"),
                  ("HUD", "MobileCastleHudRoot, HUDController", "Single Wall, minimal cycle, bottom abilities"),
              ], [1750, 3650, 3960], compact=True)
    page_break(doc)

    # 27 DATA CONTRACTS
    begin_page(doc, "Yeni Data Contract'ları", "Section 11 | Technical", "İdeal implementation, content eklemeyi kod değişikliğinden ayırır ve generated run state'i save edilebilir yapar.")
    add_table(doc,
              ["Contract", "Minimum alanlar", "Owner"],
              [
                  ("EnemyDefinition", "Id, entity prefab, base stats, pool prewarm/expand, spawn weight", "Enemy catalog"),
                  ("RunDifficultyProfile", "BaseSpawn curve, phase multipliers, active cap, backlog policy", "Difficulty SO"),
                  ("HeartNodeDefinition", "Tags, effects, rarity, depth, repeatable, cost growth, conflicts", "Node catalog"),
                  ("GeneratedRunGraph", "Seed, node ids, edges, hidden/revealed state, levels, locks", "Run save"),
                  ("WorkerAllocation", "Four target ratios, actual counts, caps, idle population", "ECS singleton"),
                  ("ArcherFormation", "40 cells, 25 local seeds/points, version", "Placement owner"),
                  ("ActiveAbilityState", "Unlock, cooldown remaining, tuning multipliers", "Run save"),
                  ("CouncilRunState", "Regular day index, flags, recent templates, active effects", "Run save"),
                  ("MetaState", "Currency, upgrade levels, pool unlocks, tutorial flags", "Meta save"),
              ], [2050, 5000, 2310], compact=True)
    add_heading(doc, "Migration ilkesi", 2)
    add_para(doc, "Mevcut TechTreeCatalog sabit node tanımlarını koruyabilir; fakat runtime graph tanım asset'lerini değiştirmez. Generator, katalogdan seçilmiş node id'leri ve edge'leri run state'e yazar. Save load aynı graph'ı yeniden kurar; source asset runtime state taşımaz.")
    add_callout(doc, "Do not rename for aesthetics", "MobileCastle* isimleri teknik borç olarak belgelenebilir; yeni PC/Steam tasarım dokümanını uygulamak için geniş rename/refactor zorunlu değildir.", "frost", "pale_blue")
    page_break(doc)

    # 28 PERFORMANCE
    begin_page(doc, "Performans Sözleşmesi", "Section 11 | Performance", "10.000 düşman ve 1.000 okçu pazarlama vaadidir; yalnız stress scene başarısı değil, gerçek oyun sırasında UI, save ve VFX ile birlikte korunmalıdır.")
    add_table(doc,
              ["Alan", "Risk", "Gerekli strateji", "Kabul"],
              [
                  ("Spawn / pool", "Structural churn", "Prewarm + expandable reusable entities", "Cap'te backlog kaybolmaz"),
                  ("Target search", "10M+ brute checks", "Spatial query + target load", "Frame spike üretmez"),
                  ("Projectiles", "1.000 high fire rate", "Pooling / burst-safe spawn / lifetime cleanup", "Ammo truth korunur"),
                  ("VFX/SFX", "Her hit feedback", "Budget, aggregation, rate limit", "Okunur mix"),
                  ("Worker visuals", "Sınırsız pop", "Representative visual density", "Gameplay count'tan ayrışmaz"),
                  ("Save", "10k entity snapshot", "Compact deterministic state", "Continue perceptually faithful"),
              ], [1450, 2300, 3500, 2110], compact=True)
    add_heading(doc, "Ölçüm senaryoları", 2)
    for text in [
        "1.000 okçu + 10.000 düşman + projectile peak + Night post-processing.",
        "Fireball radius içinde en yoğun horde ve aynı frame çoklu death return-to-pool.",
        "Arrow refill sırasında 1.000 okçunun yeniden ateşe başlaması.",
        "Ana menü save/continue sırasında maksimum run state.",
        "Düşük, orta ve yüksek worker visual density geçişi.",
    ]:
        add_bullet(doc, text, bullet_num)
    add_callout(doc, "Design over implementation", "10.000 teknik olarak sürdürülemezse ilk çözüm düşman HP şişirmek değildir. Targeting, projectile, VFX ve pool maliyetleri optimize edilir; eşzamanlı hedef değişikliği owner kararı gerektirir.", "heart", "pale_red")
    page_break(doc)

    # 29 BALANCE ANALYTICS
    begin_page(doc, "Balance Değişkenleri ve Telemetry", "Section 12 | Tuning", "Kesin sayılar tasarım gerçeği değil, ölçülebilir tuning yüzeyidir. Formüller SO/profile içinde görünür olmalıdır.")
    add_table(doc,
              ["Sistem", "Ana tuning değişkenleri", "Başarı sinyali"],
              [
                  ("Spawn", "Day curve, phase multiplier, backlog, active cap", "Run length dağılımı; cap saturation"),
                  ("Wall", "Base HP, repair cost, repair amount, Emergency %", "Night başına damage; Stone spend"),
                  ("Economy", "Base rates, capacity cost, efficiency growth", "Idle pop; resource bottleneck"),
                  ("Population", "Food per arrival, bed curve, dawn count", "Pop growth; unused beds"),
                  ("Archers", "Type base stats, cost growth, retrain cost", "Type mix; DPS; arrow drain"),
                  ("Heart", "Essence drop, node base cost, growth, rarity", "Node depth; unspent Essence"),
                  ("Council", "Regular cadence fixed; effect bands, repeat memory, decision timer", "Choice rate; expiry rate; option split; run impact"),
                  ("Meta", "Reward weights, upgrade costs/effects", "Run-to-run acceleration"),
              ], [1500, 4700, 3160], compact=True)
    add_heading(doc, "Önerilen event'ler", 2)
    add_table(doc,
              ["Event", "Payload"],
              [
                  ("run_started", "meta levels, starting resources, graph seed/version"),
                  ("phase_changed", "day index, phase, alive enemies, spawn backlog"),
                  ("resource_spent", "resource, amount, purchase type, resulting level/count"),
                  ("archer_changed", "buy/retrain, type from/to, total cap usage"),
                  ("heart_node_bought", "node id, level, depth, cost, revealed children"),
                  ("council_resolved", "day, template id, option/expired, effects, next-night delta"),
                  ("ability_cast", "ability, phase, cooldown, targets hit / repair amount"),
                  ("wall_repaired", "phase, Stone cost, HP before/after"),
                  ("run_ended", "day, kills, peak enemies, peak pop, Wall damage timeline, meta reward"),
              ], [2700, 6660], compact=True)
    page_break(doc)

    # 30 DELIVERY
    begin_page(doc, "Uygulama Sırası", "Section 13 | Production", "Bu sıra gün veya hafta taahhüdü değildir. Her paket bir sonraki paketin doğrulanabilir temelini üretir.")
    add_figure(doc, fig["delivery"], "Şekil 10. Bağımlılık odaklı uygulama sırası.", "Eight-stage implementation dependency sequence from contracts to product gate.")
    add_callout(doc, "Sequence rule", "Bir sonraki pakete geçmek için önceki paketin kabul kriterleri Play Mode veya ilgili EditMode testleriyle kanıtlanır. Görsel polish, performans ve save en sona ertelenen temizlik değildir; her pakette regresyon kapısıdır.", "green", "pale_green")
    page_break(doc)

    # 31 WORK PACKAGES 1
    begin_page(doc, "Work Packages A-D", "Section 13 | Production", "İlk dört paket koşunun veri sözleşmesini, savaşı, ekonomiyi ve 1.000 okçuluk garnizonu kurar.")
    packages = [
        ("A | System Contracts", "Tek Wall, run/meta state ayrımı, resource upkeep yasağı ve tuning owner'larını kod contract'ına sabitle.", "Save schema, config defaults, old Gate/Core paths disabled", "Run reset ve Continue deterministik"),
        ("B | Continuous Horde", "60 saniye fazlar, quantity-only growth, tek enemy catalog ve 10k-expandable pool.", "Cycle, spawn budget, backlog, return-to-pool", "Stats sabitken gün baskısı artar"),
        ("C | Economy + Population", "Dört worker oranı, hazır binalar, capacity/efficiency, yatak ve dawn arrivals.", "No passive drain; worker world representation", "Yeni pop tek Food öder; cap aşılmaz"),
        ("D | Archers + Ammo", "Üç tür, 1.000 ortak cap, instant buy/retrain, 40×25 placement, scalable targeting.", "Arrow refill, stable formation, target load", "1k×10k stress senaryosu çalışır"),
    ]
    add_table(doc, ["Paket", "Amaç", "Ana çıktı", "Kabul kapısı"], packages, [1700, 3000, 2800, 1860], compact=True)
    add_heading(doc, "A-D test özeti", 2)
    for text in [
        "Wall HP sıfırlandığında aynı frame anında ve yalnız bir kez Game Over.",
        "Enemy stats 1. gün ve ileri günlerde aynı; değişen spawn count/budget.",
        "Food 0 iken mevcut nüfus eksilmez; yeni arrival bütçeyle sınırlanır.",
        "1.001. okçu satın alınamaz; retrain toplam sayıyı değiştirmez.",
        "Save/load 40×25 formasyonu ve okçu tür sayılarını korur.",
    ]:
        add_bullet(doc, text, bullet_num)
    page_break(doc)

    # 32 WORK PACKAGES 2
    begin_page(doc, "Work Packages E-I", "Section 13 | Production", "İkinci grup procedural build, Council, aktif yetenek, meta/save ve ürün kalite kapısını tamamlar.")
    packages2 = [
        ("E | Castle Heart", "Generated graph, Essence currency, reveal, repeatable/evolution/Keystone.", "Guaranteed core nodes + saved hidden graph", "Aynı seed/load aynı graph"),
        ("F | Council", "3/6/9... regular cadence, curated context-aware choices.", "Schedule + effect guardrails + saved memory", "Takvim bozulmaz; hiçbir etki ana cap'leri bypass etmez"),
        ("G | Active Abilities", "Fireball, Rally, Emergency Repair tek cooldown barında.", "World targeting, global buff, Night repair", "Kaynak tüketmez; input UI ile çakışmaz"),
        ("H | Meta + Persistence", "Death-only reward, fixed upgrade list, single run save.", "Idempotent reward + migration guard", "Force-close ölümü geri alamaz"),
        ("I | Product Gate", "Minimal HUD, first-run tutorial, day/night art/audio, stress QA.", "PC input, single Wall bar, polish mockups", "10k scenario + clean render + tutorial pass"),
    ]
    add_table(doc, ["Paket", "Amaç", "Ana çıktı", "Kabul kapısı"], packages2, [1700, 3000, 2800, 1860], compact=True)
    add_heading(doc, "E-I test özeti", 2)
    for text in [
        "Graph hiçbir koşuda Rapid/Frost/Fireball'ı unreachable üretmez.",
        "Normal node Keystone gibi başka yolu kapatmaz.",
        "Heart açıkken cycle, spawn, worker ve cooldown tamamen durur.",
        "Council günleri 3/6/9... düzenini korur ve aynı gün yalnız bir kez açılır.",
        "Council nüfus, okçu, Wall ve horde etkileri ilgili ana sistem guardrail'lerini aşmaz.",
        "Meta upgrade aktif run graph'ını geriye dönük değiştirmez.",
        "Tutorial tamamlandıktan sonra ikinci run'da otomatik görünmez.",
    ]:
        add_bullet(doc, text, bullet_num)
    page_break(doc)

    # 33 TEST MATRIX
    begin_page(doc, "Test Matrisi", "Section 14 | QA", "Kabul testleri yalnız happy path'i değil, tasarım sözleşmesini bozabilecek sınırları hedefler.")
    add_table(doc,
              ["Alan", "Test", "Beklenen"],
              [
                  ("Cycle", "60 saniye tam loop", "30/5/20/5; kesintisiz spawn"),
                  ("Horde", "Aktif cap dolu", "Talep backlog'a gider, silinmez"),
                  ("Wall", "Night normal repair", "Buton kapalı; Stone harcanmaz"),
                  ("Wall", "HP 0 + aynı frame repair", "Game Over kazanır"),
                  ("Population", "Food yetersiz dawn", "Mevcut pop korunur; arrival sınırlanır"),
                  ("Ammo", "Arrow 0 / refill", "Ateş durur / satın alım anında başlar"),
                  ("Archers", "40 tile'da 1.000 spawn", "Her tile 25 stable point; cap aşılmaz"),
                  ("Targeting", "Yoğun overkill", "Incoming damage load hedefleri dağıtır"),
                  ("Heart", "Invalid generated graph", "Validation reroll; run başlamazsa açık hata"),
                  ("Council", "Day 3/6/9 Dawn + reload/phase replay", "Düzenli kart aynı gün yalnız bir kez gelir; takvim kaymaz"),
                  ("Council", "Pop/archer/defense/horde etkisi", "Bed+Food, 1.000 cap, Wall-only ve count-only korunur"),
                  ("Save", "Menu çıkış / Continue", "Aynı graph, phase, Wall, economy"),
                  ("Death", "Process restart after Game Over", "Meta bir kez; run sıfır"),
                  ("HUD", "16:9 / ultrawide", "Sabit kompozisyon ve kritik UI kırpılmaz"),
              ], [1550, 3600, 4210], compact=True)
    page_break(doc)

    # 34 RISK
    begin_page(doc, "Risk Register", "Section 14 | Risk", "Riskler özellik listesini büyütmek için değil, onaylı kimliği korumak için yönetilir.")
    add_table(doc,
              ["Risk", "Etki", "Erken sinyal", "Mitigation / kill rule"],
              [
                  ("1k×10k targeting", "Frame collapse", "Night spike", "Spatial query + load; HP scaling'e kaçma"),
                  ("Projectile/VFX flood", "CPU/GPU ve görsel çamur", "Hit event backlog", "Budget + aggregation + pool"),
                  ("Graph unreachable", "Run brick", "Core unlock yok", "Generator validation + deterministic fallback"),
                  ("Meta runaway", "Erken oyun anlamsız", "Day 1 instant clear", "Diminishing values; reward curve telemetry"),
                  ("Ammo tıklama angaryası", "Idle kimliği bozulur", "Refill frequency çok yüksek", "Paket/capacity/efficiency tune; auto-spend ekleme"),
                  ("HUD tekrar büyür", "Savaş görünmez", "Kalıcı paneller", "Tek drawer + sabit layout + mockup gate"),
                  ("Legacy path leakage", "Gate/Core, wave, old tech geri gelir", "Duplicate UI/state", "Source owner audit + legacy guard tests"),
                  ("Unreviewed Council content", "Generic/slop tone", "Anlamsız seçim veya tekrar", "Authored template review + effect budget test"),
                  ("Unapproved lore/content", "Owner intent drift", "Boss veya canon fraksiyon eklenir", "Out-of-scope listesi kod review kapısı"),
              ], [1950, 1800, 2300, 3310], compact=True)
    page_break(doc)

    # 35 DOD
    begin_page(doc, "Release Definition of Done", "Section 14 | Product Gate", "Blueprint ancak aşağıdaki oyuncu ve teknik sözleşmeler birlikte sağlandığında uygulanmış sayılır.")
    dod = [
        "Koşu yalnız Wall HP = 0 ile biter; final wave, boss veya ikinci fail phase yoktur.",
        "Çıkış catalog'u tek düşman prefabı içerir; difficulty enemy stats değil adet büyütür.",
        "60 saniyelik döngü doğru sürelerle kesintisiz çalışır ve speed/offline sistemi yoktur.",
        "Wood, Stone, Iron, Food pasif negatif akmaz; Arrow dışında sürekli tüketim yoktur.",
        "Nüfus, yatak ve worker oranları save/load ile korunur; worker world feedback doğru yoğunluğu gösterir.",
        "Basic/Rapid/Frost toplam 1.000 cap ve 40×25 stable placement sözleşmesine uyar.",
        "Arrow satın alma anlıktır; Fletcher / production queue yoktur.",
        "Castle Heart generated graph doğrulanır, save edilir ve yalnız Grave Essence kullanır.",
        "Council 3/6/9... günlerinde çalışır ve yalnız onaylı şablon/etki havuzunu kullanır.",
        "Council etkileri yatak+Food, 1.000 okçu cap, tek Wall ve count-only difficulty sözleşmelerini bypass etmez.",
        "Fireball, Rally ve Emergency Repair alt orta HUD'da cooldown ile çalışır.",
        "Meta ödülü yalnız ölümde bir kez verilir; gönüllü reset yoktur.",
        "HUD tek Wall barı kullanır; Cycle UI minimal ve ayrı polish onayından geçmiştir.",
        "İlk-run tutorial tamamlanır; sonraki run'da tekrar etmez.",
        "1.000 okçu + 10.000 düşman stress senaryosu hedef donanımda kabul edilen frame pacing'i sağlar.",
        "EditMode/PlayMode testleri, save migration ve uzun-run soak raporu temizdir.",
    ]
    for item in dod:
        add_bullet(doc, item, bullet_num)
    add_callout(doc, "Quality gate", "Polish ayrı bir iş paketi değil, bu DoD'nin parçasıdır. Faz göstergesi, Wall barı, day/night lighting, horde mix ve combat feedback görsel inceleme olmadan tamam sayılmaz.", "green", "pale_green")
    page_break(doc)

    # 36 OUT OF SCOPE
    begin_page(doc, "Bilinçli Kapsam Dışı", "Section 15 | Guardrails", "Aşağıdaki sistemler bu blueprint'in çıkış sözleşmesine dahil değildir. Varlıkları ancak yeni owner kararıyla yeniden değerlendirilir.")
    out = [
        ("Boss / miniboss / elite / enemy variant", "Tek prefab ve sayı tehdidi kararını bozar"),
        ("Blood Moon / sabit special-night takvimi", "Sabit special-night pacing tek prefab ve sayı tehdidi kararını bozar; ayrı isimli/periyodik mod onaylanmamıştır"),
        ("Düşman scout / forecast", "Oyuncuya anlamlı karar üretmediği için reddedildi"),
        ("Enemy lane / front selection", "Düşman saldırıyı seçer; oyuncu geleni seçmez"),
        ("Build grid / building placement", "İzometrik kompozisyon hazır ve art-directed kalır"),
        ("Fletcher production", "Ammo Wood ile anında satın alınır"),
        ("Archer death / individual HP", "Wall tek savunma truth'udur"),
        ("Gate / Core HP", "Tek Wall barı ve tek Game Over condition"),
        ("Arrow Storm", "Fireball + 1.000 okçu ile rol tekrarı"),
        ("Voluntary reset / prestige button", "Ölüm ve run riskini zayıflatır"),
        ("Offline income / offline death", "Menü ve kapalı uygulamada koşu donar"),
        ("Separate archer upgrade panel", "Castle Heart tek teknoloji owner'ıdır"),
        ("Mobile ads / rewarded / IAP", "Yeni ürün çerçevesi PC/Steam etkileşimidir"),
    ]
    add_table(doc, ["Kapsam dışı", "Neden"], out, [3300, 6060], compact=True)
    add_callout(doc, "Existing dormant content", "Moat, Blood Moon veya legacy wave kodu repoda bulunabilir. Bu blueprint bunları otomatik olarak silmez; ancak yeni core loop'a aktif bağlanmaları ayrıca onay gerektirir. Council bu listenin parçası değildir; aktif core run sistemidir.", "fire", "pale_gold")
    page_break(doc)

    # 37 OPEN POLISH
    begin_page(doc, "Açık Polish ve İçerik Kararları", "Section 15 | Open Decisions", "Ana gameplay soruları tamamdır. Aşağıdaki maddeler uygulamayı durdurmayan, fakat final sunum öncesi owner onayı isteyen creative işlerdir.")
    add_table(doc,
              ["Açık konu", "Kilitli sınır", "Gerekli çıktı"],
              [
                  ("Faz göstergesi", "Minimal; büyük DAY/DUSK/NIGHT paneli yok", "2-3 HUD mockup ve motion örneği"),
                  ("Meta para adı", "Ölüm sonrası kalıcı currency", "İsim + ikon + death screen copy"),
                  ("Narrative premise", "Heart / Essence sistemleriyle çelişmez", "World pitch + opening copy"),
                  ("Heart node pool", "Rapid/Frost/Fireball garanti", "Launch node catalog + effect specs"),
                  ("Council launch içeriği", "Free-form AI yok; yalnız curated şablon ve etkiler", "Onaylı template/atom listesi + tekrar ve bütçe testi"),
                  ("Keystone çiftleri", "Yalnız birbirini kapatır", "En az 3 onaylı trade-off çifti"),
                  ("Fireball evolution pool", "İlk büyü; mana yok", "VFX ve 2-3 evolution spec"),
                  ("Exact tuning", "Stats değil enemy count büyür", "Profile curves + telemetry targetları"),
                  ("Day/night audio", "Fazı UI'dan bağımsız anlatır", "Mix map + rate-limit budgets"),
              ], [2200, 3500, 3660], compact=True)
    add_callout(doc, "Review method", "Bu açık maddeler kod içinde varsayımla kapatılmaz. Önce mockup/spec, sonra owner kararı, sonra implementation yapılır.", "heart", "pale_red")
    page_break(doc)

    # 38 SOURCE AUDIT
    begin_page(doc, "Read-Only Source Audit", "Appendix | Evidence", "Aşağıdaki dosyalar mevcut owner'ları ve aktif repo gerçeğini anlamak için okundu. Hiçbiri bu doküman oluşturulurken değiştirilmedi.")
    sources = [
        ("CLAUDE.md + GDD v5.0 + MASTER_PLAN", "Aktif yön haritası, sürekli kuşatma ve tarihsel ürün bağlamı"),
        ("NewGameScene + MobileCastleCombatSubScene", "Kamera, outside tilemap, ECS authoring ve aktif component bağlantıları"),
        ("MobileCastleHudRoot.prefab + HUDController.cs", "Aktif HUD truth: cycle panel, savunma barları ve drawer yapısı"),
        ("MobileCastleArcherTilePlacement.cs", "40 tile, modulo/stack ve center-out yerleşim"),
        ("ArcherShootSystem.cs + ArrowHitSystem.cs", "Closest-to-archer araması, ok hasarı ve Frost slow"),
        ("ArcherComponents.cs + ArcherDefinitionSO.cs", "Basic/Rapid/Frost runtime verisi ve maliyet tanımları"),
        ("TechNodeDefinitionSO + TechTreeCatalogSO + TechTreeUI", "Node etkileri, catalog/reveal sözleşmesi ve graph layout"),
        ("CouncilComposer + CouncilEventUI + CouncilEventCatalog", "Bağlam-duyarlı deterministik seçim, UI zamanlaması, authored template/atom havuzu ve save hafızası"),
        ("GameManager.cs + MetaUpgradeSO.cs", "Okçu satın alma, spell bridge, save ve meta etkileri"),
        ("SpellCastUI.cs", "Fireball targeting ve cooldown arayüzü"),
        ("BuildingComponents.cs + ResourceComponents.cs", "Hazır binalar, üretim kaynakları ve ArrowSupply"),
    ]
    add_table(doc, ["Kaynak", "Neden incelendi"], sources, [4200, 5160], compact=True)
    add_callout(doc, "Deliberate exclusion", "DEAD_WALLS_NEW_FEATURE_MILESTONE_PLAN.docx ve .pdf yalnız eski dosyaların ayrı kaldığını doğrulamak için kontrol edildi; bu blueprint'in içerik veya tasarım kaynağı olarak kullanılmadı.", "heart", "pale_red")
    add_heading(doc, "Final authority order", 2)
    for text in [
        "1. Bu konuşmada owner tarafından kilitlenen kararlar.",
        "2. Bu blueprint içindeki LOCKED sözleşmeler.",
        "3. Aktif prefab / scene / code owner'ları.",
        "4. Eski GDD / master plan yalnız tarihsel bağlam.",
    ]:
        add_bullet(doc, text, bullet_num)

    # Final paragraph
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(20)
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run("END OF BLUEPRINT  |  BUILD THE WALL, THEN LET THE NUMBERS BECOME THE MONSTER")
    set_run_font(r, size=9, color=COLORS["heart"], bold=True)

    # Document-level widow/orphan safety and language metadata.
    for para in doc.paragraphs:
        para.paragraph_format.widow_control = True
        for run in para.runs:
            rpr = run._element.get_or_add_rPr()
            lang = rpr.find(qn("w:lang"))
            if lang is None:
                lang = OxmlElement("w:lang")
                rpr.append(lang)
            lang.set(qn("w:val"), "tr-TR")

    doc.save(OUT_DOCX)
    return OUT_DOCX


def main():
    figures = create_figures()
    path = build_document(figures)
    print(path)


if __name__ == "__main__":
    main()
